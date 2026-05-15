import os
from datetime import date
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.utils.helpers import log, add_hyperlink, add_image_border, highlight_run
from backend.utils.constants import WOS_COLLECTIONS

class ReportService:
    def __init__(self):
        pass

    def add_heading_tnr(self, doc, text, level=1):
        """Adds a heading with Times New Roman font."""
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "Times New Roman"

    def add_caption(self, doc, text):
        """Adds a small caption for images. Only 'Nota:' is italicized."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.5)
        
        if text.startswith("Nota:"):
            run_n = p.add_run("Nota:")
            run_n.font.name = "Times New Roman"
            run_n.font.size = Pt(9)
            run_n.italic = True
            
            rest = text[5:] # Length of "Nota:" is 5
            run_r = p.add_run(rest)
            run_r.font.name = "Times New Roman"
            run_r.font.size = Pt(9)
            run_r.italic = False
        else:
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.italic = False

    def generate_word_report(self, meta, pubs_by_year, total_docs, retracted_scopus,
                             predatory_hits, year_chart_buf, country_chart_buf,
                             pubs_by_country, output_file, nro_informe, institucion,
                             retracted_wos=0):
        try:
            doc = Document()
            
            # Page Setup
            section = doc.sections[0]
            section.page_width    = Cm(21)
            section.page_height   = Cm(29.7)
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(2.5)

            # Style Config
            style_normal = doc.styles["Normal"]
            style_normal.font.name = "Times New Roman"
            style_normal.font.size = Pt(11)
            style_normal.paragraph_format.line_spacing = 1.5
            style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for h_level, pt_size, indent_cm, before, after in [(1, 11, 0.0, 12, 6), (2, 11, 0.5, 6, 4), (3, 11, 1.0, 3, 2)]:
                try:
                    style = doc.styles[f"Heading {h_level}"]
                    style.font.name = "Times New Roman"; style.font.size = Pt(pt_size); style.font.bold = True
                    style.font.color.rgb = RGBColor(0, 0, 0)
                    style.paragraph_format.left_indent = Cm(indent_cm)
                    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
                    style.paragraph_format.line_spacing = 1.5; style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    rPr = style.element.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rPr.insert(0, rFonts)
                except: pass

            # ── Header ───────────────────────────────────────────────────────────
            for line in institucion:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(line); r.font.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0

            doc.add_paragraph() # Spacer

            # Informe N°
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("Informe N° IDIC-IQInteg-R-"); r.font.bold = True
            r2 = p.add_run(str(nro_informe)); r2.font.bold = True
            if str(nro_informe).startswith("["): highlight_run(r2)

            # Journal Title
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("Integridad científica de la revista "); r.font.bold = True
            rj = p.add_run(meta.get('journal', '')); rj.font.bold = True; rj.font.italic = True

            # Date
            today = date.today()
            meses = ["enero","febrero","marzo","abril","mayo","junio","julio","agosto","septiembre","octubre","noviembre","diciembre"]
            fecha_str = f"Lima, {today.day:02d} de {meses[today.month-1]} de {today.year}"
            pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.add_run(fecha_str)

            # Article Reference
            if meta.get("article_title"):
                p_hdr = doc.add_paragraph(); p_hdr.paragraph_format.space_after = Pt(6)
                r_hdr = p_hdr.add_run("Artículo de referencia analizado"); r_hdr.bold = True
                for lbl, val in [("Título", meta.get("article_title")), ("Autor(es)", meta.get("article_author")), ("DOI", meta.get("article_doi")), ("Fecha", meta.get("article_date"))]:
                    if val:
                        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
                        r_l = p.add_run(f"{lbl}: "); r_l.bold = True
                        p.add_run(str(val))

            # ── 1. Datos de la revista ───────────────────────────────────────────
            self.add_heading_tnr(doc, "1. Datos de la revista", level=1)
            
            def add_subitem(label, value, indent=0.5, spacing=4, highlight_placeholder=None):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(indent)
                p.paragraph_format.space_before = Pt(spacing); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"{label}: "); r.bold = True
                if value and not str(value).startswith("["):
                    p.add_run(str(value))
                elif highlight_placeholder:
                    highlight_run(p.add_run(highlight_placeholder))
                return p

            # 1.1 ISSN / E-ISSN
            issn = meta.get('issn_print')
            eissn = meta.get('eissn')
            p_issn = doc.add_paragraph(); p_issn.paragraph_format.left_indent = Cm(0.5)
            p_issn.paragraph_format.space_before = Pt(4); p_issn.paragraph_format.space_after = Pt(2)
            
            if issn and eissn:
                r1 = p_issn.add_run("1.1 ISSN: "); r1.bold = True
                p_issn.add_run(f"{issn};  ")
                r2 = p_issn.add_run("E-ISSN: "); r2.bold = True
                p_issn.add_run(f"{eissn}")
            elif issn:
                r1 = p_issn.add_run("1.1 ISSN: "); r1.bold = True
                p_issn.add_run(str(issn))
            elif eissn:
                r2 = p_issn.add_run("1.1 E-ISSN: "); r2.bold = True
                p_issn.add_run(str(eissn))
            else:
                r1 = p_issn.add_run("1.1 ISSN: "); r1.bold = True
                highlight_run(p_issn.add_run("[n/a]"))
            
            pweb = add_subitem("1.2 Sitio web", "")
            wh = meta.get('homepage')
            if wh and str(wh).startswith('http'): add_hyperlink(pweb, wh, wh)
            else: pweb.add_run(str(wh) if wh else '[n/a]')
            
            add_subitem("1.3 Editorial", meta.get('publisher', ''))
            add_subitem("1.4 APC", meta.get('apc'), highlight_placeholder="[COMPLETAR: monto y moneda, o 'Sin APC']")
            add_subitem("1.5 Tiempo de publicación de artículos", meta.get('pub_time'), highlight_placeholder="[COMPLETAR: días promedio]")

            # 1.6 CiteScore
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"1.6 Cuartil CiteScore (Scopus {meta.get('citescore_year', '2024')}):"); r.bold = True
            
            quartiles = meta.get("quartiles", {})
            if quartiles:
                for area, items in quartiles.items():
                    p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(0.5)
                    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
                    p2.paragraph_format.line_spacing = 1.0
                    p2.add_run("De acuerdo con el área temática ")
                    ra = p2.add_run(area); ra.italic = True
                    p2.add_run(", la revista se clasifica en las siguientes categorías y cuartiles:")
                    for it in items:
                        p_it = doc.add_paragraph()
                        p_it.paragraph_format.left_indent = Cm(1.5)
                        p_it.paragraph_format.first_line_indent = Cm(-0.5)
                        p_it.paragraph_format.space_before = Pt(0); p_it.paragraph_format.space_after = Pt(0)
                        rb = p_it.add_run("•  "); rb.bold = True; rb.font.size = Pt(14)
                        p_it.add_run(f"Se ubica en el cuartil {it['cuartil'].replace('Q','')} ")
                        rq = p_it.add_run(f"({it['cuartil']})"); rq.bold = True
                        p_it.add_run(" en la categoría ")
                        rs = p_it.add_run(it['subarea']); rs.italic = True
                        p_it.add_run(".")
            else:
                p_nc = doc.add_paragraph(); p_nc.paragraph_format.left_indent = Cm(0.5)
                highlight_run(p_nc.add_run("[Sin datos de cuartiles Scopus]"))

            # 1.7 WoS
            jcr_year = date.today().year - 2
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"1.7 Cuartil WoS (JIF {jcr_year}):"); r.bold = True
            
            wos_cats = meta.get("wos_categories", [])
            wos_cols = meta.get("wos_collections", [])
            wos_q_map = meta.get("wos_quartiles", {})
            
            if wos_cats or wos_cols:
                for cat in wos_cats:
                    cat_name = cat.get("name") if isinstance(cat, dict) else str(cat)
                    p_cat = doc.add_paragraph()
                    p_cat.paragraph_format.left_indent = Cm(1.0)
                    p_cat.paragraph_format.first_line_indent = Cm(-0.5)
                    p_cat.paragraph_format.space_before = Pt(0); p_cat.paragraph_format.space_after = Pt(0)
                    rb = p_cat.add_run("•  "); rb.bold = True; rb.font.size = Pt(14)
                    p_cat.add_run("Se ubica en el cuartil ")
                    # Check mapping for manual quartiles entered in frontend
                    q_val = wos_q_map.get(cat_name, "") if isinstance(wos_q_map, dict) else ""
                    if q_val and not str(q_val).startswith("["):
                        q_num = str(q_val).replace('Q','')
                        p_cat.add_run(f"{q_num} ")
                        rq = p_cat.add_run(f"({q_val})"); rq.bold = True
                    else:
                        highlight_run(p_cat.add_run("[Q?]"))
                    p_cat.add_run(" en la categoría ")
                    rc = p_cat.add_run(cat_name); rc.italic = True; p_cat.add_run(".")
                
                # Additional collections
                all_cols = list(wos_cols)
                wosm = str(meta.get("wos_collections_manual", "")).strip()
                if wosm:
                    for c in wosm.split(','):
                        if c.strip() and c.strip() not in all_cols: all_cols.append(c.strip())
                
                for col in all_cols:
                    p_col = doc.add_paragraph()
                    p_col.paragraph_format.left_indent = Cm(1.0)
                    p_col.paragraph_format.first_line_indent = Cm(-0.5)
                    p_col.paragraph_format.space_before = Pt(0); p_col.paragraph_format.space_after = Pt(0)
                    rb = p_col.add_run("•  "); rb.bold = True; rb.font.size = Pt(14)
                    p_col.add_run("Forma parte de la colección "); rc = p_col.add_run(str(col)); rc.italic = True; p_col.add_run(".")
            else:
                p_nw = doc.add_paragraph(); p_nw.paragraph_format.left_indent = Cm(0.5)
                highlight_run(p_nw.add_run("[COMPLETAR: cuartil y categoría WoS]"))

            # 1.8 Coverage
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_before = Pt(4)
            r = p.add_run("1.8 Vigencia en Scopus: "); r.bold = True
            p.add_run(f"Revista vigente en Scopus de {meta.get('coverage_start')} a {meta.get('coverage_end') or 'presente'}.")
            
            # Alerta Revista Descontinuada
            disc = meta.get("discontinued")
            if disc and isinstance(disc, dict) and disc.get("is_discontinued"):
                # Texto al costado (sin negrita ni rojo)
                p.add_run(" (cobertura descontinuada en Scopus)")
                
                p_disc = doc.add_paragraph(); p_disc.paragraph_format.left_indent = Cm(0.5)
                r_disc = p_disc.add_run(f"⚠ ALERTA: Esta revista figura como DESCONTINUADA en Scopus. Cobertura: {disc.get('coverage', 'n/a')}."); r_disc.bold = True
                highlight_run(r_disc)
            
            sid = meta.get("source_id")
            if sid:
                sc_link = f"https://www.scopus.com/sourceid/{sid}"
                p_lnk = doc.add_paragraph(); p_lnk.paragraph_format.left_indent = Cm(0.5)
                add_hyperlink(p_lnk, sc_link, sc_link)

            # ── 2. Producción científica ─────────────────────────────────────────
            self.add_heading_tnr(doc, "2. Producción científica de la revista", level=1)
            
            self.add_heading_tnr(doc, "2.1 Publicaciones por año", level=2)
            if total_docs > 0 or pubs_by_year:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                
                if total_docs > 0:
                    p.add_run(f"La revista {meta.get('journal')} registra, a la fecha, {total_docs:,} documentos indexados en Scopus. ")
                
                if pubs_by_year:
                    years = sorted([int(y) for y in pubs_by_year.keys()])
                    if years:
                        y_start, y_end = years[0], years[-1]
                        c_start = pubs_by_year.get(str(y_start), 0)
                        vals = list(pubs_by_year.values())
                        c_max = max(vals) if vals else 0
                        y_max = max(pubs_by_year, key=pubs_by_year.get)
                        p.add_run(f"Durante el período {y_start}–{y_end}, el volumen inició con {c_start:,} documentos en {y_start} y alcanzó su punto más alto en {y_max} con {c_max:,} documentos.")

            if year_chart_buf:
                p_fig = doc.add_paragraph(); p_fig.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fig.paragraph_format.left_indent = Cm(0.5)
                p_fig.add_run("Figura 1").bold = True
                doc.add_paragraph().add_run("Publicaciones por año de la revista").italic = True
                doc.paragraphs[-1].paragraph_format.left_indent = Cm(0.5)
                # Figura 1: Chart
                p_img_para = doc.add_picture(year_chart_buf, width=Cm(14))
                p_img = doc.paragraphs[-1]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # line_spacing=1.15 y space_before de 12pt mejoran la visibilidad de los bordes
                p_img.paragraph_format.line_spacing = 1.15
                p_img.paragraph_format.space_before = Pt(12)
                p_img.paragraph_format.space_after = Pt(0)
                p_img.paragraph_format.left_indent = Cm(0) 
                
                add_image_border(doc)
                self.add_caption(doc, "Nota: Elaborado a partir de datos extraídos de Scopus.")

            self.add_heading_tnr(doc, "2.2 Documentos por país", level=2)
            if pubs_by_country:
                # Mapping for country translation to Spanish
                country_map = {
                    "United States": "Estados Unidos", "China": "China", "United Kingdom": "Reino Unido",
                    "Germany": "Alemania", "Japan": "Japón", "France": "Francia", "Canada": "Canadá",
                    "Australia": "Australia", "Italy": "Italia", "Spain": "España", "South Korea": "Corea del Sur",
                    "India": "India", "Netherlands": "Países Bajos", "Brazil": "Brasil", "Switzerland": "Suiza",
                    "Sweden": "Suecia", "Poland": "Polonia", "Turkey": "Turquía", "Iran": "Irán",
                    "Portugal": "Portugal", "Belgium": "Bélgica", "Denmark": "Dinamarca", "Norway": "Noruega",
                    "Finland": "Finlandia", "Austria": "Austria", "Czech Republic": "República Checa",
                    "Greece": "Grecia", "Israel": "Israel", "New Zealand": "Nueva Zelanda", "Singapore": "Singapur",
                    "Argentina": "Argentina", "Mexico": "México", "South Africa": "Sudáfrica", "Russia": "Rusia",
                    "Taiwan": "Taiwán", "Malaysia": "Malasia", "Thailand": "Tailandia", "Egypt": "Egipto",
                    "Saudi Arabia": "Arabia Saudita", "Nigeria": "Nigeria", "Colombia": "Colombia",
                    "Chile": "Chile", "Indonesia": "Indonesia", "Pakistan": "Pakistán", "Romania": "Rumania",
                    "Hungary": "Hungría", "Slovakia": "Eslovaquia", "Croatia": "Croacia", "Serbia": "Serbia",
                    "Ukraine": "Ucrania", "Peru": "Perú", "Vietnam": "Vietnam", "Morocco": "Marruecos",
                    "Tunisia": "Túnez", "Kenya": "Kenia"
                }

                top_10 = pubs_by_country[:10]
                names = []
                for i, d in enumerate(top_10):
                    country_en = d["country"]
                    country_es = country_map.get(country_en, country_en)
                    
                    if i == 0:
                        name = f"{country_es} ({d['count']:,} documentos)"
                    else:
                        name = country_es
                    names.append(name)

                if len(names) > 1:
                    last = names[-1]
                    rest = names[:-1]
                    prefix = ", ".join(rest)
                    separator = " e " if last.lower().startswith(("i", "hi")) and not last.lower().startswith("hie") else " y "
                    top_str = f"{prefix}{separator}{last}"
                else:
                    top_str = names[0] if names else ""

                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                p.add_run(f"A nivel de concentración de documentos por país, destacan los aportes de autores de {top_str}.")
            
            if country_chart_buf:
                p_fig2 = doc.add_paragraph(); p_fig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fig2.paragraph_format.left_indent = Cm(0.5)
                p_fig2.add_run("Figura 2").bold = True
                doc.add_paragraph().add_run("Publicaciones de la revista por país").italic = True
                doc.paragraphs[-1].paragraph_format.left_indent = Cm(0.5)
                # Figura 2: Chart
                p_img_para2 = doc.add_picture(country_chart_buf, width=Cm(14))
                p_img2 = doc.paragraphs[-1]
                p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # line_spacing=1.15 y space_before de 12pt mejoran la visibilidad de los bordes
                p_img2.paragraph_format.line_spacing = 1.15
                p_img2.paragraph_format.space_before = Pt(12)
                p_img2.paragraph_format.space_after = Pt(0)
                p_img2.paragraph_format.left_indent = Cm(0)

                add_image_border(doc)
                self.add_caption(doc, "Nota: Elaborado a partir de datos extraídos de Scopus.")

            self.add_heading_tnr(doc, "2.3 Artículos retractados por la revista", level=2)
            
            # Scopus
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            rb1 = p.add_run("•  "); rb1.bold = True; rb1.font.size = Pt(14)
            if retracted_scopus == 0:
                p.add_run("No se identificaron artículos retractados en Scopus.")
            elif retracted_scopus == 1:
                p.add_run("Se identificó un 1 artículo retractado en Scopus.")
            else:
                p.add_run(f"Se identificaron {retracted_scopus} artículos retractados en Scopus.")
            
            # Web of Science
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.0)
            p2.paragraph_format.first_line_indent = Cm(-0.5)
            p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
            rb2 = p2.add_run("•  "); rb2.bold = True; rb2.font.size = Pt(14)
            if retracted_wos is not None and retracted_wos >= 0:
                if retracted_wos == 0:
                    p2.add_run("No se identificaron artículos retractados en Web of Science.")
                elif retracted_wos == 1:
                    p2.add_run("Se identificó un 1 artículo retractado en Web of Science.")
                else:
                    p2.add_run(f"Se identificaron {retracted_wos} artículos retractados en Web of Science.")
            else:
                highlight_run(p2.add_run("[COMPLETAR: retractados en WoS]"))

            # Retraction Watch (CSV Data)
            rw = meta.get("retracted_watch")
            if rw and rw.get("count", 0) > 0:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Cm(1.0)
                p3.paragraph_format.first_line_indent = Cm(-0.5)
                p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(0)
                rb3 = p3.add_run("•  "); rb3.bold = True; rb3.font.size = Pt(14)
                p3.add_run(f"La base de datos Retraction Watch identifica {rw['count']} registros adicionales.")
                
                # Motivos principales
                if rw.get("top_reasons"):
                    p_reasons = doc.add_paragraph()
                    p_reasons.paragraph_format.left_indent = Cm(1.5)
                    p_reasons.add_run(f"Motivos principales: {', '.join(rw['top_reasons'])}.")

                # Ejemplos (Máximo 3)
                p_ex_title = doc.add_paragraph()
                p_ex_title.paragraph_format.left_indent = Cm(1.5)
                p_ex_title.add_run("Ejemplos recientes:").italic = True
                
                for ex in rw.get("examples", []):
                    pe = doc.add_paragraph()
                    pe.paragraph_format.left_indent = Cm(2.0)
                    pe.paragraph_format.space_before = Pt(0); pe.paragraph_format.space_after = Pt(0)
                    pe.add_run(f"- {ex['title']} ({ex['date']})").font.size = Pt(9)
                    pe_r = doc.add_paragraph()
                    pe_r.paragraph_format.left_indent = Cm(2.5)
                    pe_r.paragraph_format.space_before = Pt(0); pe_r.paragraph_format.space_after = Pt(0)
                    pe_r.add_run(f"Motivo: {ex['reasons']}").font.size = Pt(8); pe_r.italic = True

            # ── 3. Listas predatorias ────────────────────────────────────────────
            self.add_heading_tnr(doc, "3. Verificación de listas de revistas predatorias", level=1)
            
            lista1_url = "https://www.predatoryjournals.org/the-list/journals"
            lista2_url = "https://beallslist.net/"

            if not predatory_hits:
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                p.add_run("No figura en la lista de Predatory journals ni en Beall's List.")
            else:
                # Agrupar hits por categoría (Revista vs Editorial) y Fuente (Lista 1 vs Lista 2)
                j_sources = []
                p_sources = []
                
                for hit in predatory_hits:
                    source_name = ""
                    if "Lista 1" in hit: source_name = "Predatory Journals"
                    elif "Lista 2" in hit: source_name = "Beall's List"
                    
                    if hit.startswith("Revista"):
                        if source_name and source_name not in j_sources: j_sources.append(source_name)
                    elif hit.startswith("Editorial"):
                        if source_name and source_name not in p_sources: p_sources.append(source_name)

                # --- Lógica de Redacción Estilizada ---
                def get_src_str(sources):
                    if not sources: return ""
                    plural = "s" if len(sources) > 1 else ""
                    return f"lista{plural} de {' y '.join(sources)}"

                sentence = ""
                if j_sources and p_sources:
                    if set(j_sources) == set(p_sources):
                        # Escenario: Ambas en lo mismo
                        sentence = f"La revista y su editorial figuran en la {get_src_str(j_sources)}."
                    else:
                        # Escenario: Coincidencia Mixta
                        txt_j = get_src_str(j_sources)
                        txt_p = get_src_str(p_sources).replace("listas de ", "").replace("lista de ", "")
                        sentence = f"La revista figura en la {txt_j}, y su editorial en la de {txt_p}."
                elif j_sources:
                    sentence = f"La revista figura en la {get_src_str(j_sources)}."
                elif p_sources:
                    sentence = f"El editorial de la revista figura en la {get_src_str(p_sources)}."

                if sentence:
                    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                    r = p.add_run(sentence); r.bold = True

            bp1 = doc.add_paragraph()
            bp1.paragraph_format.left_indent = Cm(1.0)
            bp1.paragraph_format.first_line_indent = Cm(-0.5)
            bp1.paragraph_format.space_before = Pt(0); bp1.paragraph_format.space_after = Pt(0)
            rb1 = bp1.add_run("•  "); rb1.bold = True; rb1.font.size = Pt(14)
            add_hyperlink(bp1, lista1_url, lista1_url)
            
            bp2 = doc.add_paragraph()
            bp2.paragraph_format.left_indent = Cm(1.0)
            bp2.paragraph_format.first_line_indent = Cm(-0.5)
            bp2.paragraph_format.space_before = Pt(0); bp2.paragraph_format.space_after = Pt(0)
            rb2 = bp2.add_run("•  "); rb2.bold = True; rb2.font.size = Pt(14)
            add_hyperlink(bp2, lista2_url, lista2_url)

            # ── 4. Conclusiones ──────────────────────────────────────────────────
            self.add_heading_tnr(doc, "4. Conclusiones", level=1)
            p1 = doc.add_paragraph(); p1.paragraph_format.left_indent = Cm(0.5); p1.add_run("4.1 ").bold = True
            p1.add_run(f"Es una revista indexada en Scopus (vigente de {meta.get('coverage_start')} a {meta.get('coverage_end') or 'presente'}).")
            
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(0.5); p2.add_run("4.2 ").bold = True
            best_q = "N/A"
            if quartiles:
                all_q = [it["cuartil"] for items in quartiles.values() for it in items]
                best_q = sorted(set(all_q))[0] if all_q else "N/A"
            p2.add_run(f"Presenta métricas CiteScore con clasificación en {best_q} en sus categorías de mayor desempeño.")
            
            p3 = doc.add_paragraph(); p3.paragraph_format.left_indent = Cm(0.5); p3.add_run("4.3 ").bold = True
            if not predatory_hits:
                p3.add_run("No presenta indicios de malas prácticas (no figura en listas predatorias).")
            else:
                r = p3.add_run("Figura en listas de revistas predatorias. Se recomienda cautela."); r.font.color.rgb = RGBColor(192, 0, 0)

            # 4.4 Discontinued Conclusion
            disc = meta.get("discontinued")
            if disc and isinstance(disc, dict) and disc.get("is_discontinued"):
                p4 = doc.add_paragraph(); p4.paragraph_format.left_indent = Cm(0.5); p4.add_run("4.4 ").bold = True
                r4 = p4.add_run("La revista figura como DESCONTINUADA por Scopus. Se recomienda verificar la idoneidad de la publicación."); r4.font.color.rgb = RGBColor(192, 0, 0); r4.bold = True

            # ── 5. Recomendaciones ───────────────────────────────────────────────
            self.add_heading_tnr(doc, "5. Recomendaciones", level=1)
            p51 = doc.add_paragraph(); p51.paragraph_format.left_indent = Cm(0.5); p51.add_run("5.1 ").bold = True
            p51.add_run("Para elecciones futuras, se recomienda leer el ")
            PROTOCOLO_URL = "https://www.ulima.edu.pe/sites/default/files/centers/files/protocolo_idoneidad_publicaciones_idic_v25jul_1.pdf"
            add_hyperlink(p51, PROTOCOLO_URL, "Protocolo para la evaluación de la idoneidad de publicaciones científicas")
            p51.add_run(". Esta revisión permitirá comprender mejor los criterios de calidad que la institución considera importantes al evaluar las publicaciones científicas.")
            
            p52 = doc.add_paragraph(); p52.paragraph_format.left_indent = Cm(0.5); p52.add_run("5.2 ").bold = True
            p52.add_run("Es importante reconocer que el comportamiento editorial, la calidad y la reputación de una revista pueden cambiar con el tiempo. Por lo tanto, se aconseja verificar sus indexaciones periódicamente.")

            # Footer
            footer = doc.sections[0].footer
            pf = footer.paragraphs[0]; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rf = pf.add_run("El presente informe tiene una vigencia de tres (03) meses a partir de su fecha de emisión."); rf.font.size = Pt(9); rf.italic = True

            doc.save(output_file)
            return output_file
        except Exception as e:
            log(f"ERROR:Fallo masivo en generate_word_report: {e}")
            raise e
