"""
Generador de Informe de Integridad Científica - Backend
Genera informe Word (.docx) con datos de Scopus API
"""

import sys
import re
import io
import time
import json
import requests
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from io import StringIO, BytesIO
from datetime import date
from bs4 import BeautifulSoup
import os
import concurrent.futures

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── ASJC MAP (abreviado) ──────────────────────────────────────────────────────
ASJC = {
    # ── Multidisciplinary ─────────────────────────────────────────────────────
    "10": ("Multidisciplinary", "Multidisciplinary"),
    "1000": ("Multidisciplinary", "Multidisciplinary"),

    # ── Agricultural and Biological Sciences ─────────────────────────────────
    "11": ("Agricultural and Biological Sciences", "General"),
    "1100": ("Agricultural and Biological Sciences", "General Agricultural and Biological Sciences"),
    "1101": ("Agricultural and Biological Sciences", "Agricultural and Biological Sciences (miscellaneous)"),
    "1102": ("Agricultural and Biological Sciences", "Agronomy and Crop Science"),
    "1103": ("Agricultural and Biological Sciences", "Animal Science and Zoology"),
    "1104": ("Agricultural and Biological Sciences", "Aquatic Science"),
    "1105": ("Agricultural and Biological Sciences", "Ecology, Evolution, Behavior and Systematics"),
    "1106": ("Agricultural and Biological Sciences", "Food Science"),
    "1107": ("Agricultural and Biological Sciences", "Forestry"),
    "1108": ("Agricultural and Biological Sciences", "Horticulture"),
    "1109": ("Agricultural and Biological Sciences", "Insect Science"),
    "1110": ("Agricultural and Biological Sciences", "Plant Science"),
    "1111": ("Agricultural and Biological Sciences", "Soil Science"),

    # ── Arts and Humanities ──────────────────────────────────────────────────
    "12": ("Arts and Humanities", "General"),
    "1200": ("Arts and Humanities", "General Arts and Humanities"),
    "1201": ("Arts and Humanities", "Arts and Humanities (miscellaneous)"),
    "1202": ("Arts and Humanities", "History"),
    "1203": ("Arts and Humanities", "Language and Linguistics"),
    "1204": ("Arts and Humanities", "Archaeology"),
    "1205": ("Arts and Humanities", "Classics"),
    "1206": ("Arts and Humanities", "Conservation"),
    "1207": ("Arts and Humanities", "History and Philosophy of Science"),
    "1208": ("Arts and Humanities", "Literature and Literary Theory"),
    "1209": ("Arts and Humanities", "Museology"),
    "1210": ("Arts and Humanities", "Music"),
    "1211": ("Arts and Humanities", "Philosophy"),
    "1212": ("Arts and Humanities", "Religious studies"),
    "1213": ("Arts and Humanities", "Visual Arts and Performing Arts"),

    # ── Biochemistry, Genetics and Molecular Biology ──────────────────────────
    "13": ("Biochemistry, Genetics and Molecular Biology", "General"),
    "1300": ("Biochemistry, Genetics and Molecular Biology", "General Biochemistry, Genetics and Molecular Biology"),
    "1301": ("Biochemistry, Genetics and Molecular Biology", "Biochemistry, Genetics and Molecular Biology (miscellaneous)"),
    "1302": ("Biochemistry, Genetics and Molecular Biology", "Ageing"),
    "1303": ("Biochemistry, Genetics and Molecular Biology", "Biochemistry"),
    "1304": ("Biochemistry, Genetics and Molecular Biology", "Biophysics"),
    "1305": ("Biochemistry, Genetics and Molecular Biology", "Biotechnology"),
    "1306": ("Biochemistry, Genetics and Molecular Biology", "Cancer Research"),
    "1307": ("Biochemistry, Genetics and Molecular Biology", "Cell Biology"),
    "1308": ("Biochemistry, Genetics and Molecular Biology", "Clinical Biochemistry"),
    "1309": ("Biochemistry, Genetics and Molecular Biology", "Developmental Biology"),
    "1310": ("Biochemistry, Genetics and Molecular Biology", "Endocrinology"),
    "1311": ("Biochemistry, Genetics and Molecular Biology", "Genetics"),
    "1312": ("Biochemistry, Genetics and Molecular Biology", "Molecular Biology"),
    "1313": ("Biochemistry, Genetics and Molecular Biology", "Molecular Medicine"),
    "1314": ("Biochemistry, Genetics and Molecular Biology", "Physiology"),
    "1315": ("Biochemistry, Genetics and Molecular Biology", "Structural Biology"),

    # ── Business, Management and Accounting ──────────────────────────────────
    "14": ("Business, Management and Accounting", "General"),
    "1400": ("Business, Management and Accounting", "General Business, Management and Accounting"),
    "1401": ("Business, Management and Accounting", "Business, Management and Accounting (miscellaneous)"),
    "1402": ("Business, Management and Accounting", "Accounting"),
    "1403": ("Business, Management and Accounting", "Business and International Management"),
    "1404": ("Business, Management and Accounting", "Management Information Systems"),
    "1405": ("Business, Management and Accounting", "Management of Technology and Innovation"),
    "1406": ("Business, Management and Accounting", "Marketing"),
    "1407": ("Business, Management and Accounting", "Organizational Behavior and Human Resource Management"),
    "1408": ("Business, Management and Accounting", "Strategy and Management"),
    "1409": ("Business, Management and Accounting", "Tourism, Leisure and Hospitality Management"),
    "1410": ("Business, Management and Accounting", "Industrial Relations"),

    # ── Chemical Engineering ──────────────────────────────────────────────────
    "15": ("Chemical Engineering", "General"),
    "1500": ("Chemical Engineering", "General Chemical Engineering"),
    "1501": ("Chemical Engineering", "Chemical Engineering (miscellaneous)"),
    "1502": ("Chemical Engineering", "Bioengineering"),
    "1503": ("Chemical Engineering", "Catalysis"),
    "1504": ("Chemical Engineering", "Chemical Health and Safety"),
    "1505": ("Chemical Engineering", "Colloid and Surface Chemistry"),
    "1506": ("Chemical Engineering", "Filtration and Separation"),
    "1507": ("Chemical Engineering", "Fluid Flow and Transfer Processes"),
    "1508": ("Chemical Engineering", "Process Chemistry and Technology"),

    # ── Chemistry ──────────────────────────────────────────────────────────
    "16": ("Chemistry", "General"),
    "1600": ("Chemistry", "General Chemistry"),
    "1601": ("Chemistry", "Chemistry (miscellaneous)"),
    "1602": ("Chemistry", "Analytical Chemistry"),
    "1603": ("Chemistry", "Electrochemistry"),
    "1604": ("Chemistry", "Inorganic Chemistry"),
    "1605": ("Chemistry", "Organic Chemistry"),
    "1606": ("Chemistry", "Physical and Theoretical Chemistry"),
    "1607": ("Chemistry", "Spectroscopy"),

    # ── Computer Science ──────────────────────────────────────────────────────
    "17": ("Computer Science", "General"),
    "1700": ("Computer Science", "General Computer Science"),
    "1701": ("Computer Science", "Computer Science (miscellaneous)"),
    "1702": ("Computer Science", "Artificial Intelligence"),
    "1703": ("Computer Science", "Computational Theory and Mathematics"),
    "1704": ("Computer Science", "Computer Graphics and Computer-Aided Design"),
    "1705": ("Computer Science", "Computer Networks and Communications"),
    "1706": ("Computer Science", "Computer Science Applications"),
    "1707": ("Computer Science", "Computer Vision and Pattern Recognition"),
    "1708": ("Computer Science", "Hardware and Architecture"),
    "1709": ("Computer Science", "Human-Computer Interaction"),
    "1710": ("Computer Science", "Information Systems"),
    "1711": ("Computer Science", "Signal Processing"),
    "1712": ("Computer Science", "Software"),

    # ── Decision Sciences ──────────────────────────────────────────────────────
    "18": ("Decision Sciences", "General"),
    "1800": ("Decision Sciences", "General Decision Sciences"),
    "1801": ("Decision Sciences", "Decision Sciences (miscellaneous)"),
    "1802": ("Decision Sciences", "Information Systems and Management"),
    "1803": ("Decision Sciences", "Management Science and Operations Research"),
    "1804": ("Decision Sciences", "Statistics, Probability and Uncertainty"),

    # ── Earth and Planetary Sciences ──────────────────────────────────────────
    "19": ("Earth and Planetary Sciences", "General"),
    "1900": ("Earth and Planetary Sciences", "General Earth and Planetary Sciences"),
    "1901": ("Earth and Planetary Sciences", "Earth and Planetary Sciences (miscellaneous)"),
    "1902": ("Earth and Planetary Sciences", "Atmospheric Science"),
    "1903": ("Earth and Planetary Sciences", "Computers in Earth Sciences"),
    "1904": ("Earth and Planetary Sciences", "Earth-Surface Processes"),
    "1905": ("Earth and Planetary Sciences", "Economic Geology"),
    "1906": ("Earth and Planetary Sciences", "Geochemistry and Petrology"),
    "1907": ("Earth and Planetary Sciences", "Geology"),
    "1908": ("Earth and Planetary Sciences", "Geophysics"),
    "1909": ("Earth and Planetary Sciences", "Geotechnical Engineering and Engineering Geology"),
    "1910": ("Earth and Planetary Sciences", "Oceanography"),
    "1911": ("Earth and Planetary Sciences", "Paleontology"),
    "1912": ("Earth and Planetary Sciences", "Space and Planetary Science"),
    "1913": ("Earth and Planetary Sciences", "Stratigraphy"),

    # ── Economics, Econometrics and Finance ──────────────────────────────────
    "20": ("Economics, Econometrics and Finance", "General"),
    "2000": ("Economics, Econometrics and Finance", "General Economics, Econometrics and Finance"),
    "2001": ("Economics, Econometrics and Finance", "Economics, Econometrics and Finance (miscellaneous)"),
    "2002": ("Economics, Econometrics and Finance", "Economics and Econometrics"),
    "2003": ("Economics, Econometrics and Finance", "Finance"),

    # ── Energy ──────────────────────────────────────────────────────────────
    "21": ("Energy", "General"),
    "2100": ("Energy", "General Energy"),
    "2101": ("Energy", "Energy (miscellaneous)"),
    "2102": ("Energy", "Energy Engineering and Power Technology"),
    "2103": ("Energy", "Fuel Technology"),
    "2104": ("Energy", "Nuclear Energy and Engineering"),
    "2105": ("Energy", "Renewable Energy, Sustainability and the Environment"),

    # ── Engineering ──────────────────────────────────────────────────────────
    "22": ("Engineering", "General"),
    "2200": ("Engineering", "General Engineering"),
    "2201": ("Engineering", "Engineering (miscellaneous)"),
    "2202": ("Engineering", "Aerospace Engineering"),
    "2203": ("Engineering", "Automotive Engineering"),
    "2204": ("Engineering", "Biomedical Engineering"),
    "2205": ("Engineering", "Civil and Structural Engineering"),
    "2206": ("Engineering", "Computational Mechanics"),
    "2207": ("Engineering", "Control and Systems Engineering"),
    "2208": ("Engineering", "Electrical and Electronic Engineering"),
    "2209": ("Engineering", "Industrial and Manufacturing Engineering"),
    "2210": ("Engineering", "Mechanical Engineering"),
    "2211": ("Engineering", "Mechanics of Materials"),
    "2212": ("Engineering", "Ocean Engineering"),
    "2213": ("Engineering", "Safety, Risk, Reliability and Quality"),
    "2214": ("Engineering", "Media Technology"),
    "2215": ("Engineering", "Building and Construction"),
    "2216": ("Engineering", "Architecture"),

    # ── Environmental Science ──────────────────────────────────────────────────
    "23": ("Environmental Science", "General"),
    "2300": ("Environmental Science", "General Environmental Science"),
    "2301": ("Environmental Science", "Environmental Science (miscellaneous)"),
    "2302": ("Environmental Science", "Ecological Modeling"),
    "2303": ("Environmental Science", "Ecology"),
    "2304": ("Environmental Science", "Environmental Chemistry"),
    "2305": ("Environmental Science", "Environmental Engineering"),
    "2306": ("Environmental Science", "Global and Planetary Change"),
    "2307": ("Environmental Science", "Health, Toxicology and Mutagenesis"),
    "2308": ("Environmental Science", "Management, Monitoring, Policy and Law"),
    "2309": ("Environmental Science", "Nature and Landscape Conservation"),
    "2310": ("Environmental Science", "Pollution"),
    "2311": ("Environmental Science", "Waste Management and Disposal"),
    "2312": ("Environmental Science", "Water Science and Technology"),

    # ── Immunology and Microbiology ──────────────────────────────────────────
    "24": ("Immunology and Microbiology", "General"),
    "2400": ("Immunology and Microbiology", "General Immunology and Microbiology"),
    "2401": ("Immunology and Microbiology", "Immunology and Microbiology (miscellaneous)"),
    "2402": ("Immunology and Microbiology", "Applied Microbiology and Biotechnology"),
    "2403": ("Immunology and Microbiology", "Immunology"),
    "2404": ("Immunology and Microbiology", "Microbiology"),
    "2405": ("Immunology and Microbiology", "Parasitology"),
    "2406": ("Immunology and Microbiology", "Virology"),

    # ── Materials Science ──────────────────────────────────────────────────────
    "25": ("Materials Science", "General"),
    "2500": ("Materials Science", "General Materials Science"),
    "2501": ("Materials Science", "Materials Science (miscellaneous)"),
    "2502": ("Materials Science", "Biomaterials"),
    "2503": ("Materials Science", "Ceramics and Composites"),
    "2504": ("Materials Science", "Electronic, Optical and Magnetic Materials"),
    "2505": ("Materials Science", "Materials Chemistry"),
    "2506": ("Materials Science", "Metals and Alloys"),
    "2507": ("Materials Science", "Polymers and Plastics"),
    "2508": ("Materials Science", "Surfaces, Coatings and Films"),

    # ── Mathematics ──────────────────────────────────────────────────────────
    "26": ("Mathematics", "General"),
    "2600": ("Mathematics", "General Mathematics"),
    "2601": ("Mathematics", "Mathematics (miscellaneous)"),
    "2602": ("Mathematics", "Algebra and Number Theory"),
    "2603": ("Mathematics", "Analysis"),
    "2604": ("Mathematics", "Applied Mathematics"),
    "2605": ("Mathematics", "Computational Mathematics"),
    "2606": ("Mathematics", "Control and Optimization"),
    "2607": ("Mathematics", "Discrete Mathematics and Combinatorics"),
    "2608": ("Mathematics", "Geometry and Topology"),
    "2609": ("Mathematics", "Logic"),
    "2610": ("Mathematics", "Mathematical Physics"),
    "2611": ("Mathematics", "Modelling and Simulation"),
    "2612": ("Mathematics", "Numerical Analysis"),
    "2613": ("Mathematics", "Statistics and Probability"),
    "2614": ("Mathematics", "Theoretical Computer Science"),

    # ── Medicine ──────────────────────────────────────────────────────────────
    "27": ("Medicine", "General"),
    "2700": ("Medicine", "General Medicine"),
    "2701": ("Medicine", "Medicine (miscellaneous)"),
    "2702": ("Medicine", "Anatomy"),
    "2703": ("Medicine", "Anesthesiology and Pain Medicine"),
    "2704": ("Medicine", "Biochemistry, medical"),
    "2705": ("Medicine", "Cardiology and Cardiovascular Medicine"),
    "2706": ("Medicine", "Critical Care and Intensive Care Medicine"),
    "2707": ("Medicine", "Complementary and alternative medicine"),
    "2708": ("Medicine", "Dermatology"),
    "2709": ("Medicine", "Drug guides"),
    "2710": ("Medicine", "Embryology"),
    "2711": ("Medicine", "Emergency Medicine"),
    "2712": ("Medicine", "Endocrinology, Diabetes and Metabolism"),
    "2713": ("Medicine", "Epidemiology"),
    "2714": ("Medicine", "Family Practice"),
    "2715": ("Medicine", "Gastroenterology"),
    "2716": ("Medicine", "Genetics(clinical)"),
    "2717": ("Medicine", "Geriatrics and Gerontology"),
    "2718": ("Medicine", "Health Informatics"),
    "2719": ("Medicine", "Health Policy"),
    "2720": ("Medicine", "Hematology"),
    "2721": ("Medicine", "Hepatology"),
    "2722": ("Medicine", "Histology"),
    "2723": ("Medicine", "Immunology and Allergy"),
    "2724": ("Medicine", "Internal Medicine"),
    "2725": ("Medicine", "Infectious Diseases"),
    "2726": ("Medicine", "Microbiology (medical)"),
    "2727": ("Medicine", "Nephrology"),
    "2728": ("Medicine", "Clinical Neurology"),
    "2729": ("Medicine", "Obstetrics and Gynaecology"),
    "2730": ("Medicine", "Oncology"),
    "2731": ("Medicine", "Ophthalmology"),
    "2732": ("Medicine", "Orthopedics and Sports Medicine"),
    "2733": ("Medicine", "Otorhinolaryngology"),
    "2734": ("Medicine", "Pathology and Forensic Medicine"),
    "2735": ("Medicine", "Pediatrics, Perinatology, and Child Health"),
    "2736": ("Medicine", "Pharmacology (medical)"),
    "2737": ("Medicine", "Physiology (medical)"),
    "2738": ("Medicine", "Psychiatry and Mental health"),
    "2739": ("Medicine", "Public Health, Environmental and Occupational Health"),
    "2740": ("Medicine", "Pulmonary and Respiratory Medicine"),
    "2741": ("Medicine", "Radiology Nuclear Medicine and imaging"),
    "2742": ("Medicine", "Rehabilitation"),
    "2743": ("Medicine", "Reproductive Medicine"),
    "2744": ("Medicine", "Reviews and References, Medical"),
    "2745": ("Medicine", "Rheumatology"),
    "2746": ("Medicine", "Surgery"),
    "2747": ("Medicine", "Transplantation"),
    "2748": ("Medicine", "Urology"),

    # ── Neuroscience ──────────────────────────────────────────────────────────
    "28": ("Neuroscience", "General"),
    "2800": ("Neuroscience", "General Neuroscience"),
    "2801": ("Neuroscience", "Neuroscience (miscellaneous)"),
    "2802": ("Neuroscience", "Behavioral Neuroscience"),
    "2803": ("Neuroscience", "Biological Psychiatry"),
    "2804": ("Neuroscience", "Cellular and Molecular Neuroscience"),
    "2805": ("Neuroscience", "Cognitive Neuroscience"),
    "2806": ("Neuroscience", "Developmental Neuroscience"),
    "2807": ("Neuroscience", "Endocrine and Autonomic Systems"),
    "2808": ("Neuroscience", "Neurology"),
    "2809": ("Neuroscience", "Sensory Systems"),

    # ── Nursing ──────────────────────────────────────────────────────────────
    "29": ("Nursing", "General"),
    "2900": ("Nursing", "General Nursing"),
    "2901": ("Nursing", "Nursing (miscellaneous)"),
    "2902": ("Nursing", "Advanced and Specialised Nursing"),
    "2903": ("Nursing", "Assessment and Diagnosis"),
    "2904": ("Nursing", "Care Planning"),
    "2905": ("Nursing", "Community and Home Care"),
    "2906": ("Nursing", "Critical Care"),
    "2907": ("Nursing", "Emergency"),
    "2908": ("Nursing", "Fundamentals and skills"),
    "2909": ("Nursing", "Gerontology"),
    "2910": ("Nursing", "Issues, ethics and legal aspects"),
    "2911": ("Nursing", "Leadership and Management"),
    "2912": ("Nursing", "LPN and LVN"),
    "2913": ("Nursing", "Maternity and Midwifery"),
    "2914": ("Nursing", "Medical-Surgical"),
    "2915": ("Nursing", "Nurse Assisting"),
    "2916": ("Nursing", "Nutrition and Dietetics"),
    "2917": ("Nursing", "Oncology (nursing)"),
    "2918": ("Nursing", "Pathophysiology"),
    "2919": ("Nursing", "Pediatrics"),
    "2920": ("Nursing", "Pharmacology (nursing)"),
    "2921": ("Nursing", "Psychiatric Mental Health"),
    "2922": ("Nursing", "Research and Theory"),
    "2923": ("Nursing", "Review and Exam Preparation"),

    # ── Pharmacology, Toxicology and Pharmaceutics ──────────────────────────
    "30": ("Pharmacology, Toxicology and Pharmaceutics", "General"),
    "3000": ("Pharmacology, Toxicology and Pharmaceutics", "General Pharmacology, Toxicology and Pharmaceutics"),
    "3001": ("Pharmacology, Toxicology and Pharmaceutics", "Pharmacology, Toxicology and Pharmaceutics (miscellaneous)"),
    "3002": ("Pharmacology, Toxicology and Pharmaceutics", "Drug Discovery"),
    "3003": ("Pharmacology, Toxicology and Pharmaceutics", "Pharmaceutical Science"),
    "3004": ("Pharmacology, Toxicology and Pharmaceutics", "Pharmacology"),
    "3005": ("Pharmacology, Toxicology and Pharmaceutics", "Toxicology"),

    # ── Physics and Astronomy ──────────────────────────────────────────────────
    "31": ("Physics and Astronomy", "General"),
    "3100": ("Physics and Astronomy", "General Physics and Astronomy"),
    "3101": ("Physics and Astronomy", "Physics and Astronomy (miscellaneous)"),
    "3102": ("Physics and Astronomy", "Acoustics and Ultrasonics"),
    "3103": ("Physics and Astronomy", "Astronomy and Astrophysics"),
    "3104": ("Physics and Astronomy", "Condensed Matter Physics"),
    "3105": ("Physics and Astronomy", "Instrumentation"),
    "3106": ("Physics and Astronomy", "Nuclear and High Energy Physics"),
    "3107": ("Physics and Astronomy", "Atomic and Molecular Physics, and Optics"),
    "3108": ("Physics and Astronomy", "Radiation"),
    "3109": ("Physics and Astronomy", "Statistical and Nonlinear Physics"),
    "3110": ("Physics and Astronomy", "Surfaces and Interfaces"),

    # ── Psychology ──────────────────────────────────────────────────────────
    "32": ("Psychology", "General"),
    "3200": ("Psychology", "General Psychology"),
    "3201": ("Psychology", "Psychology (miscellaneous)"),
    "3202": ("Psychology", "Applied Psychology"),
    "3203": ("Psychology", "Clinical Psychology"),
    "3204": ("Psychology", "Developmental and Educational Psychology"),
    "3205": ("Psychology", "Experimental and Cognitive Psychology"),
    "3206": ("Psychology", "Neuropsychology and Physiological Psychology"),
    "3207": ("Psychology", "Social Psychology"),

    # ── Social Sciences ──────────────────────────────────────────────────────
    "33": ("Social Sciences", "General"),
    "3300": ("Social Sciences", "General Social Sciences"),
    "3301": ("Social Sciences", "Social Sciences (miscellaneous)"),
    "3302": ("Social Sciences", "Archaeology"),
    "3303": ("Social Sciences", "Development"),
    "3304": ("Social Sciences", "Education"),
    "3305": ("Social Sciences", "Geography, Planning and Development"),
    "3306": ("Social Sciences", "Health (social science)"),
    "3307": ("Social Sciences", "Human Factors and Ergonomics"),
    "3308": ("Social Sciences", "Law"),
    "3309": ("Social Sciences", "Library and Information Sciences"),
    "3310": ("Social Sciences", "Linguistics and Language"),
    "3311": ("Social Sciences", "Safety Research"),
    "3312": ("Social Sciences", "Sociology and Political Science"),
    "3313": ("Social Sciences", "Transportation"),
    "3314": ("Social Sciences", "Anthropology"),
    "3315": ("Social Sciences", "Communication"),
    "3316": ("Social Sciences", "Cultural Studies"),
    "3317": ("Social Sciences", "Demography"),
    "3318": ("Social Sciences", "Gender Studies"),
    "3319": ("Social Sciences", "Life-span and Life-course Studies"),
    "3320": ("Social Sciences", "Political Science and International Relations"),
    "3321": ("Social Sciences", "Public Administration"),
    "3322": ("Social Sciences", "Urban Studies"),

    # ── Veterinary ──────────────────────────────────────────────────────────
    "34": ("Veterinary", "General"),
    "3400": ("Veterinary", "General Veterinary"),
    "3401": ("Veterinary", "Veterinary (miscellaneous)"),
    "3402": ("Veterinary", "Equine"),
    "3403": ("Veterinary", "Food Animals"),
    "3404": ("Veterinary", "Small Animals"),

    # ── Dentistry ──────────────────────────────────────────────────────────
    "35": ("Dentistry", "General"),
    "3500": ("Dentistry", "General Dentistry"),
    "3501": ("Dentistry", "Dentistry (miscellaneous)"),
    "3502": ("Dentistry", "Dental Assisting"),
    "3503": ("Dentistry", "Dental Hygiene"),
    "3504": ("Dentistry", "Oral Surgery"),
    "3505": ("Dentistry", "Orthodontics"),
    "3506": ("Dentistry", "Periodontics"),

    # ── Health Professions ──────────────────────────────────────────────────
    "36": ("Health Professions", "General"),
    "3600": ("Health Professions", "General Health Professions"),
    "3601": ("Health Professions", "Health Professions (miscellaneous)"),
    "3602": ("Health Professions", "Chiropractics"),
    "3603": ("Health Professions", "Complementary and Manual Therapy"),
    "3604": ("Health Professions", "Emergency Medical Services"),
    "3605": ("Health Professions", "Health Information Management"),
    "3606": ("Health Professions", "Medical Assisting and Transcription"),
    "3607": ("Health Professions", "Medical Laboratory Technology"),
    "3608": ("Health Professions", "Medical Terminology"),
    "3609": ("Health Professions", "Occupational Therapy"),
    "3610": ("Health Professions", "Optometry"),
    "3611": ("Health Professions", "Pharmacy"),
    "3612": ("Health Professions", "Physical Therapy, Sports Therapy and Rehabilitation"),
    "3613": ("Health Professions", "Podiatry"),
    "3614": ("Health Professions", "Radiological and Ultrasound Technology"),
    "3615": ("Health Professions", "Respiratory Care"),
    "3616": ("Health Professions", "Speech and Hearing"),
}

def get_asjc_names(code_str):
    code = str(code_str).strip()
    if code in ASJC:
        return ASJC[code]
    prefix = code[:2]
    if prefix in ASJC:
        return (ASJC[prefix][0], f"Subárea {code}")
    return (f"Código {code}", f"Subárea {code}")

def percentile_to_quartile(pct_str):
    try:
        p = float(pct_str)
        if p >= 75: return "Q1"
        if p >= 50: return "Q2"
        if p >= 25: return "Q3"
        return "Q4"
    except:
        return "N/A"

def log(msg):
    try:
        print(msg, flush=True)
    except UnicodeEncodeError:
        print(msg.encode('ascii', 'replace').decode('ascii'), flush=True)

def get_article_and_journal_from_scopus_id(api_key, scopus_id):
    """
    Dado un Scopus ID (ej. '2-s2.0-85216070591'), consulta la Abstract Retrieval API
    y devuelve un dict con metadatos del artículo y el eISSN / ISSN de su revista,
    más el source-id de Scopus para poder luego llamar a get_journal_metadata.
    """
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    # Limpiar el ID: aceptar con o sin prefijo "2-s2.0-"
    clean_id = scopus_id.strip()

    url = f"https://api.elsevier.com/content/abstract/scopus_id/{clean_id}"
    params = {"field": "dc:title,prism:publicationName,prism:issn,prism:eIssn,dc:creator,prism:coverDate,source-id,prism:doi,authkeywords"}
    resp = requests.get(url, headers=headers, params=params, timeout=20)

    if resp.status_code == 404:
        raise RuntimeError(f"Artículo con Scopus ID '{clean_id}' no encontrado")
    if resp.status_code == 401:
        raise RuntimeError("API Key inválida o sin permisos")
    if resp.status_code != 200:
        raise RuntimeError(f"Error HTTP {resp.status_code}: {resp.text[:300]}")

    data = resp.json()
    abs_resp = data.get("abstracts-retrieval-response", {})
    coredata = abs_resp.get("coredata", {})

    if not coredata:
        raise RuntimeError("Respuesta vacía de la API de Scopus para este ID")

    article_title   = coredata.get("dc:title", "N/A")
    journal_name    = coredata.get("prism:publicationName", "N/A")
    issn_print      = coredata.get("prism:issn", "")
    eissn           = coredata.get("prism:eIssn", "")
    doi             = coredata.get("prism:doi", "")
    cover_date      = coredata.get("prism:coverDate", "")
    source_id_raw   = coredata.get("source-id", "")
    creator         = coredata.get("dc:creator", "")

    # Normalizar ISSNs (Scopus a veces los devuelve con guión, a veces sin)
    def fmt_issn(s):
        s = str(s).strip().replace("-", "")
        return f"{s[:4]}-{s[4:]}" if len(s) == 8 else s

    eissn_fmt      = fmt_issn(eissn)      if eissn      else ""
    issn_print_fmt = fmt_issn(issn_print) if issn_print else ""

    # Usar eISSN si existe, si no ISSN impreso para buscar metadatos de la revista
    issn_for_journal = eissn_fmt or issn_print_fmt

    if not issn_for_journal and not source_id_raw:
        raise RuntimeError(
            f"El artículo '{article_title}' no tiene ISSN/eISSN ni source-id disponible en Scopus"
        )

    return {
        "article_title":  article_title,
        "article_doi":    doi,
        "article_date":   cover_date,
        "article_author": creator,
        "journal_name":   journal_name,
        "issn_print":     issn_print_fmt,
        "eissn":          eissn_fmt,
        "source_id":      str(source_id_raw),
        "issn_for_lookup": issn_for_journal,
    }


def get_journal_metadata(api_key, eissn):
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    url = f"https://api.elsevier.com/content/serial/title/issn/{eissn}"
    params = {"view": "CITESCORE"}
    resp = requests.get(url, headers=headers, params=params, timeout=20)
    if resp.status_code == 404:
        raise RuntimeError("ISSN no encontrado en Scopus")
    if resp.status_code == 401:
        raise RuntimeError("API Key inválida o sin permisos")
    if resp.status_code != 200:
        raise RuntimeError(f"Error HTTP {resp.status_code}: {resp.text[:200]}")

    data = resp.json().get("serial-metadata-response", {})
    entry = data.get("entry", [{}])[0]

    if "error" in entry:
        raise RuntimeError(f"Scopus error: {entry['error']}")

    homepage = scopus_link = ""
    for link in entry.get("link", []):
        ref = link.get("@ref", "")
        if ref == "homepage": homepage = link.get("@href", "")
        elif ref == "scopus-source": scopus_link = link.get("@href", "")

    cs_block = entry.get("citeScoreYearInfoList", {})
    cs_value = cs_block.get("citeScoreCurrentMetric", "")
    cs_year  = cs_block.get("citeScoreCurrentMetricYear", "")

    cs_year_list = cs_block.get("citeScoreYearInfo", [])
    if isinstance(cs_year_list, dict):
        cs_year_list = [cs_year_list]

    target = None
    for y in cs_year_list:
        if y.get("@year") == cs_year and y.get("@status") == "Complete":
            target = y; break
    if not target:
        for y in cs_year_list:
            if y.get("@status") == "Complete":
                target = y; break

    subject_ranks = []
    if target:
        try:
            subject_ranks = (
                target["citeScoreInformationList"][0]
                      ["citeScoreInfo"][0]
                      ["citeScoreSubjectRank"]
            )
            if isinstance(subject_ranks, dict):
                subject_ranks = [subject_ranks]
        except (KeyError, IndexError, TypeError):
            subject_ranks = []

    quartiles_by_area = {}
    for r in subject_ranks:
        code = r.get("subjectCode", "")
        pct  = r.get("percentile", "")
        cat, subarea = get_asjc_names(code)
        if cat not in quartiles_by_area:
            quartiles_by_area[cat] = []
        quartiles_by_area[cat].append({
            "subarea":   subarea,
            "cuartil":   percentile_to_quartile(pct),
            "percentil": pct,
        })

    src_id = entry.get("source-id", "")
    scopus_link_fmt = f"https://www.scopus.com/sourceid/{src_id}" if src_id else scopus_link

    return {
        "journal":         entry.get("dc:title", "N/A"),
        "issn_print":      entry.get("prism:issn", ""),
        "eissn":           entry.get("prism:eIssn", eissn),
        "publisher":       entry.get("dc:publisher", "N/A"),
        "homepage":        homepage,
        "scopus_link":     scopus_link_fmt,
        "coverage_start":  entry.get("coverageStartYear", "N/A"),
        "coverage_end":    entry.get("coverageEndYear", "present"),
        "citescore_value": cs_value,
        "citescore_year":  cs_year,
        "quartiles":       quartiles_by_area,
        "source_id":       src_id,
    }

def get_retracted_count(api_key, source_id):
    if not source_id: return 0
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    url = "https://api.elsevier.com/content/search/scopus"
    params = {"query": f'SOURCE-ID({source_id}) AND DOCTYPE("tb")', "count": 1, "field": "dc:title"}
    try:
        resp = requests.get(url, headers=headers, params=params, timeout=15)
        if resp.status_code == 200:
            return int(resp.json().get("search-results", {}).get("opensearch:totalResults", 0))
    except: pass
    return 0

def get_publications_by_year(api_key, source_id, start_year, end_year):
    """
    Obtiene el conteo de publicaciones por año en paralelo usando consultas explícitas.
    Garantiza resultados incluso si las facetas no están disponibles.
    """
    if not source_id: return {}
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    url = "https://api.elsevier.com/content/search/scopus"
    
    results = {}

    def fetch_year(year):
        query = f"SOURCE-ID({source_id}) AND PUBYEAR = {year}"
        params = {"query": query, "count": 1, "field": "dc:title"}
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=15)
            if resp.status_code == 200:
                total = int(resp.json().get("search-results", {}).get("opensearch:totalResults", 0))
                return (year, total) if total > 0 else None
        except:
            pass
        return None

    years = range(start_year, end_year + 1)
    with concurrent.futures.ThreadPoolExecutor(max_workers=len(years)) as executor:
        futures = [executor.submit(fetch_year, y) for y in years]
        for future in concurrent.futures.as_completed(futures):
            res = future.result()
            if res:
                y, count = res
                results[y] = count

    return results

def get_total_documents(api_key, source_id):
    if not source_id: return 0
    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    url = "https://api.elsevier.com/content/search/scopus"
    params = {"query": f"SOURCE-ID({source_id})", "count": 1, "field": "dc:title"}
    try:
        resp = requests.get(url, headers=headers, params=params, timeout=15)
        if resp.status_code == 200:
            return int(resp.json().get("search-results", {}).get("opensearch:totalResults", 0))
    except: pass
    return 0

# Lista de países más comunes en publicaciones científicas (para iterar)
COUNTRIES_TO_CHECK = [
    "United States", "China", "United Kingdom", "Germany", "Japan",
    "France", "Canada", "Australia", "Italy", "Spain",
    "South Korea", "India", "Netherlands", "Brazil", "Switzerland",
    "Sweden", "Poland", "Turkey", "Iran", "Portugal",
    "Belgium", "Denmark", "Norway", "Finland", "Austria",
    "Czech Republic", "Greece", "Israel", "New Zealand", "Singapore",
    "Argentina", "Mexico", "South Africa", "Russia", "Taiwan",
    "Malaysia", "Thailand", "Egypt", "Saudi Arabia", "Nigeria",
    "Colombia", "Chile", "Indonesia", "Pakistan", "Romania",
    "Hungary", "Slovakia", "Croatia", "Serbia", "Ukraine",
    "Peru", "Vietnam", "Morocco", "Tunisia", "Kenya",
]

def get_publications_by_country(api_key, source_id, top_n=15):
    """
    Obtiene el conteo de publicaciones por país en paralelo.
    Usa el formato de consulta LIMIT-TO (AFFILCOUNTRY, "País") recomendado.
    """
    if not source_id:
        return []

    headers = {"X-ELS-APIKey": api_key, "Accept": "application/json"}
    url     = "https://api.elsevier.com/content/search/scopus"
    
    log(f"INFO:Consultando publicaciones por país en paralelo — SOURCE-ID {source_id} ...")
    
    results = []

    def fetch_country(country_name):
        # Formato de consulta sugerido por el usuario para máxima fiabilidad
        query = f'SOURCE-ID({source_id}) AND AFFILCOUNTRY("{country_name}")'
        params = {"query": query, "count": 1, "field": "dc:title"}
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=10)
            if resp.status_code == 200:
                total = int(resp.json().get("search-results", {}).get("opensearch:totalResults", 0))
                if total > 0:
                    return {"country": country_name, "count": total}
        except:
            pass
        return None

    # Usamos un pull de hilos para procesar los países en paralelo
    # COUNTRIES_TO_CHECK tiene aprox 56 países. Con 15 hilos tomará muy poco.
    with concurrent.futures.ThreadPoolExecutor(max_workers=15) as executor:
        futures = [executor.submit(fetch_country, c) for c in COUNTRIES_TO_CHECK]
        for future in concurrent.futures.as_completed(futures):
            res = future.result()
            if res:
                results.append(res)
                log(f"INFO:  {res['country']}: {res['count']:,}")

    # Ordenar y limitar
    results.sort(key=lambda x: x["count"], reverse=True)
    return results[:top_n]


# ══════════════════════════════════════════════════════════════════════════════
#  WoS STARTER API — Colecciones, categorías JCR y retractados
# ══════════════════════════════════════════════════════════════════════════════

WOS_BASE = "https://api.clarivate.com/apis/wos-starter/v1"

# Variable global — se configura en __main__ desde params["wos_api_key"]
API_KEY_WOS = ""
WOS_HEADERS = {}

WOS_EDITION_NAMES = {
    "SCI":  "Science Citation Index (SCI)",
    "SCIE": "Science Citation Index Expanded (SCIE)",
    "SSCI": "Social Sciences Citation Index (SSCI)",
    "AHCI": "Arts & Humanities Citation Index (AHCI)",
    "ESCI": "Emerging Sources Citation Index (ESCI)",
    "BKCI": "Book Citation Index (BKCI)",
    "CPCI": "Conference Proceedings Citation Index (CPCI)",
}

WOS_COLLECTIONS = {
    "WOS.SCI":  "Science Citation Index Expanded (SCIE)",
    "WOS.SSCI": "Social Sciences Citation Index (SSCI)",
    "WOS.AHCI": "Arts & Humanities Citation Index (AHCI)",
    "WOS.ESCI": "Emerging Sources Citation Index (ESCI)",
    "WOS.ISTP": "Conference Proceedings Citation Index (CPCI)",
}


def _init_wos_api(wos_api_key: str):
    """Initialize global WoS API key and headers."""
    global API_KEY_WOS, WOS_HEADERS
    API_KEY_WOS = (wos_api_key or "").strip()
    if API_KEY_WOS:
        WOS_HEADERS = {"X-ApiKey": API_KEY_WOS}
    else:
        WOS_HEADERS = {}


def get_wos_collections(eissn: str, issn_print: str = "") -> list:
    """
    Detecta en qué colecciones WoS está indexada la revista usando hilos para acelerar la búsqueda.
    """
    if not API_KEY_WOS:
        return []

    issn_clean = (eissn or issn_print or "").replace("-", "")
    if not issn_clean:
        return []

    found = []

    def check_collection(edn_code, edn_name):
        query = f'IS={issn_clean} AND EDN==("{edn_code}")'
        try:
            resp = requests.get(
                f"{WOS_BASE}/documents",
                headers=WOS_HEADERS,
                params={"q": query, "limit": 1},
                timeout=10,
            )
            if resp.status_code == 200:
                total = resp.json().get("metadata", {}).get("total", 0)
                if total > 0:
                    return edn_name
        except:
            pass
        return None

    # Ejecutar en paralelo (son solo 5 colecciones)
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(check_collection, code, name) for code, name in WOS_COLLECTIONS.items()]
        for future in concurrent.futures.as_completed(futures):
            res = future.result()
            if res:
                found.append(res)
                log(f"INFO:  WoS coleccion encontrada: {res}")

    return found


# Lista de categorías JCR para búsqueda bruta en WoS Starter API
JCR_CATEGORIES = [
    "Acoustics","Agricultural Economics & Policy","Agricultural Engineering",
    "Agriculture Dairy & Animal Science","Agriculture Multidisciplinary",
    "Agronomy","Allergy","Anatomy & Morphology","Andrology",
    "Anesthesiology","Anthropology","Area Studies",
    "Astronomy & Astrophysics","Automation & Control Systems",
    "Behavioral Sciences","Biochemical Research Methods",
    "Biochemistry & Molecular Biology","Biodiversity Conservation",
    "Biology","Biophysics","Biotechnology & Applied Microbiology",
    "Cardiac & Cardiovascular Systems","Cell Biology",
    "Chemistry Analytical","Chemistry Applied",
    "Chemistry Inorganic & Nuclear","Chemistry Medicinal",
    "Chemistry Multidisciplinary","Chemistry Organic",
    "Chemistry Physical","Clinical Neurology","Communication",
    "Computer Science Artificial Intelligence",
    "Computer Science Cybernetics",
    "Computer Science Hardware & Architecture",
    "Computer Science Information Systems",
    "Computer Science Interdisciplinary Applications",
    "Computer Science Software Engineering",
    "Computer Science Theory & Methods","Construction & Building Technology",
    "Critical Care Medicine","Crystallography","Dentistry Oral Surgery & Medicine",
    "Dermatology & Venerereal Diseases","Developmental Biology",
    "Ecology","Economics","Education & Educational Research",
    "Education Scientific Disciplines","Electrochemistry",
    "Emergency Medicine","Endocrinology & Metabolism",
    "Energy & Fuels","Engineering Aerospace",
    "Engineering Biomedical","Engineering Chemical",
    "Engineering Environmental","Engineering Geological",
    "Engineering Industrial","Engineering Manufacturing",
    "Engineering Marine","Engineering Mechanical",
    "Engineering Multidisciplinary","Engineering Ocean",
    "Engineering Petroleum","Entomology","Environmental Sciences",
    "Environmental Studies","Ergonomics","Evolutionary Biology",
    "Fisheries","Food Science & Technology","Forestry",
    "Gastroenterology & Hepatology","Genetics & Heredity",
    "Geochemistry & Geophysics","Geography","Geography Physical",
    "Geology","Geosciences Multidisciplinary","Geriatrics & Gerontology",
    "Gerontology","Green & Sustainable Science & Technology",
    "Health Care Sciences & Services","Health Policy & Services",
    "Hematology","History & Philosophy of Science",
    "Horticulture","Hospitality Leisure Sport & Tourism",
    "Imaging Science & Photographic Technology","Immunology",
    "Infectious Diseases","Information Science & Library Science",
    "Instruments & Instrumentation","Integrative & Complementary Medicine",
    "International Relations","Marine & Freshwater Biology",
    "Materials Science Biomaterials","Materials Science Ceramics",
    "Materials Science Characterization & Testing",
    "Materials Science Coatings & Films","Materials Science Composites",
    "Materials Science Multidisciplinary","Materials Science Paper & Wood",
    "Materials Science Textiles","Mathematical & Computational Biology",
    "Mathematics","Mathematics Applied","Mathematics Interdisciplinary Applications",
    "Mechanics","Medical Ethics","Medical Informatics",
    "Medical Laboratory Technology","Medicine General & Internal",
    "Medicine Legal","Medicine Research & Experimental",
    "Metallurgy & Metallurgical Engineering","Meteorology & Atmospheric Sciences",
    "Microbiology","Microscopy","Mineralogy","Mining & Mineral Processing",
    "Multidisciplinary Sciences","Mycology","Nanoscience & Nanotechnology",
    "Neuroimaging","Neurosciences","Nuclear Science & Technology",
    "Nursing","Nutrition & Dietetics","Obstetrics & Gynecology",
    "Oceanography","Oncology","Operations Research & Management Science",
    "Ophthalmology","Optics","Orthopedics","Otorhinolaryngology",
    "Paleontology","Parasitology","Pathology","Pediatrics",
    "Peripheral Vascular Disease","Pharmacology & Pharmacy",
    "Physics Applied","Physics Atomic Molecular & Chemical",
    "Physics Condensed Matter","Physics Fluids & Plasmas",
    "Physics Mathematical","Physics Multidisciplinary",
    "Physics Nuclear","Physics Particles & Fields",
    "Physiology","Plant Sciences","Polymer Science",
    "Primary Health Care","Psychiatry","Psychology",
    "Psychology Applied","Psychology Biological",
    "Psychology Clinical","Psychology Developmental",
    "Psychology Educational","Psychology Experimental",
    "Psychology Mathematical","Psychology Multidisciplinary",
    "Psychology Social","Public Administration",
    "Public Environmental & Occupational Health",
    "Radiology Nuclear Medicine & Medical Imaging",
    "Rehabilitation","Remote Sensing","Reproductive Biology",
    "Respiratory System","Rheumatology","Robotics",
    "Social Sciences Biomedical","Social Sciences Interdisciplinary",
    "Social Sciences Mathematical Methods","Sociology",
    "Soil Science","Spectroscopy","Sport Sciences",
    "Statistics & Probability","Substance Abuse",
    "Surgery","Telecommunications","Toxicology",
    "Transplantation","Transportation","Tropical Medicine",
    "Urology & Nephrology","Veterinary Sciences","Virology",
    "Water Resources","Zoology",
]

def get_wos_categories(eissn: str, issn_print: str = "") -> list:
    """
    Detecta las categorías JCR de la revista en WoS buscando en paralelo.
    """
    if not API_KEY_WOS:
        return []

    issn_clean = (eissn or issn_print or "").replace("-", "")
    if not issn_clean:
        return []

    found = []
    
    def check_category(cat_name):
        query = f'IS={issn_clean} AND TASCA==("{cat_name}")'
        try:
            resp = requests.get(
                f"{WOS_BASE}/documents",
                headers=WOS_HEADERS,
                params={"q": query, "limit": 1},
                timeout=5,
            )
            if resp.status_code == 200:
                total = resp.json().get("metadata", {}).get("total", 0)
                if total > 0:
                    return cat_name
        except:
            pass
        return None

    log(f"INFO:Buscando categorias JCR en paralelo (mucho más rápido)...")
    
    # Procesamos en hilos con max_workers=10 para no saturar el API
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        # Enviamos los trabajos
        future_to_cat = {executor.submit(check_category, cat): cat for cat in JCR_CATEGORIES}
        
        for future in concurrent.futures.as_completed(future_to_cat):
            res = future.result()
            if res:
                found.append(res)
                log(f"INFO:  Categoria JCR encontrada: {res}")
                # Si ya encontramos 3, cancelamos el resto si es posible y salimos
                if len(found) >= 3:
                    # Nota: shutdown(wait=False) no cancela hilos ya en ejecución 
                    # pero as_completed nos permite ignorar el resto.
                    break
                    
    return found


def get_wos_retracted_count(eissn: str, issn_print: str = "") -> int:
    """
    Cuenta artículos retractados en WoS.
    Query: IS={issn} AND DT=Retracted Publication
    Retorna entero >= 0, o -1 si no se pudo consultar.
    """
    if not API_KEY_WOS:
        return -1

    for issn_val in filter(None, [eissn, issn_print]):
        issn_clean = issn_val.replace("-", "")
        try:
            resp = requests.get(
                f"{WOS_BASE}/documents",
                headers=WOS_HEADERS,
                params={
                    "q":     f"IS={issn_clean} AND DT=Retracted Publication",
                    "limit": 1,
                },
                timeout=15,
            )
            if resp.status_code == 200:
                data  = resp.json()
                total = (data.get("metadata", {}).get("total")
                         or data.get("total", 0))
                return int(total)
        except requests.RequestException:
            continue

    return -1



# ══════════════════════════════════════════════════════════════════════════════
#  SCRAPING WoS CON PLAYWRIGHT (JCR + retractados)
#  Se ejecuta solo si playwright está disponible y WoS es accesible.
# ══════════════════════════════════════════════════════════════════════════════

def _playwright_available():
    try:
        from playwright.sync_api import sync_playwright
        return True
    except ImportError:
        return False


def scrape_jcr_playwright(issn, timeout_ms=60000):
    """
    Abre jcr.clarivate.com con Playwright (Chromium headless), busca la revista
    por ISSN y extrae: JIF, cuartiles por categoría, colección WoS.

    URL de perfil directo:
      https://jcr.clarivate.com/jcr-jp/journal-profile?journal=<ISSN>&year=<año>

    Devuelve dict con la misma estructura que get_wos_journal_data, o None si falla.
    """
    if not _playwright_available():
        log("WARN:Playwright no disponible — instala: pip install playwright && playwright install chromium")
        return None

    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout

    clean   = issn.replace("-", "").strip()
    issn_fmt = f"{clean[:4]}-{clean[4:]}" if len(clean) == 8 else issn.strip()
    jcr_year = str(date.today().year - 2)

    log(f"INFO:Playwright → JCR para ISSN {issn_fmt} (año {jcr_year}) ...")

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True, args=["--no-sandbox","--disable-dev-shm-usage"])
        ctx  = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            viewport={"width": 1280, "height": 800},
        )
        page = ctx.new_page()

        result = {"found": False, "source": "jcr_playwright", "jif": "",
                  "jif_year": jcr_year, "collections": [], "categories": []}

        try:
            # ── 1. Ir al perfil directo del journal en JCR ───────────────────
            profile_url = (
                f"https://jcr.clarivate.com/jcr-jp/journal-profile"
                f"?journal={issn_fmt}&year={jcr_year}&fromPage=homejr&edition=SCIE"
            )
            page.goto(profile_url, timeout=timeout_ms, wait_until="networkidle")

            # Si redirige al login, intentar con la búsqueda general
            if "login" in page.url.lower() or "sign-in" in page.url.lower():
                log("WARN:JCR requiere login — usando búsqueda alternativa")
                browser.close()
                return None

            # Esperar que carguen los datos del journal
            try:
                page.wait_for_selector("[data-ta='journal-title'], .journal-name, h1.title, "
                                       ".jif-value, [class*='impact'], [class*='quartile']",
                                       timeout=20000)
            except PWTimeout:
                log("WARN:JCR no cargó elementos esperados — puede necesitar sesión activa")

            html = page.content()
            soup = BeautifulSoup(html, "html.parser")
            page_text = soup.get_text(" ", strip=True)

            # ── 2. Extraer JIF ───────────────────────────────────────────────
            # Patrones comunes en el HTML de JCR
            jif_patterns = [
                r"Journal Impact Factor[^\d]{0,30}(\d+\.\d+)",
                r"JIF[^\d]{0,20}(\d+\.\d+)",
                r"Impact Factor[^\d]{0,20}(\d+\.\d+)",
            ]
            for pat in jif_patterns:
                m = re.search(pat, page_text, re.IGNORECASE)
                if m:
                    result["jif"] = m.group(1)
                    break

            # También buscar en elementos específicos
            for sel in ["[data-ta='jif-value']", ".jif-value", "[class*='impact-factor'] span",
                        "[class*='JIF'] span", "app-jif-tile span"]:
                el = soup.select_one(sel)
                if el:
                    txt = el.get_text(strip=True)
                    if re.match(r"^\d+\.\d+$", txt):
                        result["jif"] = txt; break

            # ── 3. Extraer colecciones ────────────────────────────────────────
            for col in ("SCIE", "SSCI", "ESCI", "AHCI"):
                if col in page_text:
                    result["collections"].append(col)

            # ── 4. Extraer cuartiles por categoría ───────────────────────────
            # Buscar tablas con Q1/Q2/Q3/Q4
            for row in soup.select("tr"):
                cells = [td.get_text(strip=True) for td in row.select("td, th")]
                for i, cell in enumerate(cells):
                    if re.match(r"^Q[1-4]$", cell):
                        cat_name  = cells[i-1].strip() if i > 0 else ""
                        rank_str  = cells[i+1].strip() if i+1 < len(cells) else ""
                        if cat_name and len(cat_name) > 3:
                            result["categories"].append({
                                "name":     cat_name,
                                "quartile": cell,
                                "rank":     rank_str,
                            })

            # También buscar en spans/divs con "Q1","Q2","Q3","Q4"
            if not result["categories"]:
                for el in soup.select("[class*='quartile'], [data-ta*='quartile']"):
                    q_text = el.get_text(strip=True)
                    if re.match(r"^Q[1-4]$", q_text):
                        # Intentar obtener el nombre de categoría del elemento padre
                        parent = el.find_parent()
                        if parent:
                            cat_candidates = parent.get_text(strip=True)
                            cat_candidates = re.sub(r"Q[1-4]", "", cat_candidates).strip()
                            if cat_candidates and len(cat_candidates) > 3:
                                result["categories"].append({
                                    "name": cat_candidates[:80],
                                    "quartile": q_text,
                                    "rank": "",
                                })

            result["found"] = bool(result["jif"] or result["collections"] or result["categories"])

        except PWTimeout:
            log("WARN:Playwright timeout en JCR")
        except Exception as e:
            log(f"WARN:Playwright JCR error: {e}")
        finally:
            browser.close()

    return result if result["found"] else None


def scrape_wos_retracted_playwright(journal_name, issn, wos_url="https://www.webofscience.com",
                                    timeout_ms=60000):
    """
    Busca artículos retractados de la revista en Web of Science usando
    búsqueda avanzada:
      SO=("JOURNAL NAME") AND DT=(Retraction)

    Devuelve int con el número de retractados, o None si falla.
    """
    if not _playwright_available():
        return None

    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout

    log(f"INFO:Playwright -> WoS retractados para ISSN {issn} / {journal_name[:40]} ...")

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True, args=["--no-sandbox","--disable-dev-shm-usage"])
        ctx  = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                       "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            viewport={"width": 1280, "height": 800},
        )
        page = ctx.new_page()
        retracted_count = None

        try:
            adv_url = f"{wos_url}/wos/woscc/advanced-search"
            page.goto(adv_url, timeout=timeout_ms, wait_until="networkidle")

            # Detectar login
            if "login" in page.url.lower() or "sign-in" in page.url.lower():
                log("WARN:WoS requiere autenticación — retractados WoS quedará manual")
                browser.close()
                return None

            # Esperar campo de búsqueda avanzada
            try:
                page.wait_for_selector("textarea, input[placeholder*='query'], "
                                       "input[placeholder*='search'], #advancedSearchInputArea",
                                       timeout=15000)
            except PWTimeout:
                log("WARN:WoS no cargó el campo de búsqueda")
                browser.close()
                return None

            # Construir query — usar ISSN es más robusto que el nombre
            clean_issn = issn.replace("-","").strip()
            issn_fmt   = f"{clean_issn[:4]}-{clean_issn[4:]}" if len(clean_issn)==8 else issn
            query      = f'SO=("{journal_name}") AND DT=(Retraction)'

            # Limpiar campo y escribir query
            textarea = page.query_selector("textarea, #advancedSearchInputArea")
            if not textarea:
                textarea = page.query_selector("input[type='text']")
            if textarea:
                textarea.triple_click()
                textarea.fill(query)
            else:
                log("WARN:No se encontró campo de búsqueda en WoS")
                browser.close()
                return None

            # Enviar búsqueda
            page.keyboard.press("Enter")

            # Esperar resultados — buscar el número total
            try:
                page.wait_for_selector(
                    "[class*='results-count'], [class*='recordCount'], "
                    "[data-ta='results-count'], span.brand-blue",
                    timeout=25000
                )
            except PWTimeout:
                pass

            html  = page.content()
            soup  = BeautifulSoup(html, "html.parser")
            ptext = soup.get_text(" ", strip=True)

            # Buscar el número de resultados
            m = re.search(r"(\d[\d,\.]*)[\s\xa0]*(result|record|article|item|document)",
                          ptext, re.IGNORECASE)
            if m:
                num_str = m.group(1).replace(",","").replace(".","")
                retracted_count = int(num_str)
                log(f"INFO:WoS retractados encontrados: {retracted_count}")
            else:
                # Fallback: buscar "0 results" o "No results"
                if re.search(r"no results|0 results|sin resultados", ptext, re.IGNORECASE):
                    retracted_count = 0
                    log("INFO:WoS retractados: 0")

        except PWTimeout:
            log("WARN:Playwright timeout en WoS retractados")
        except Exception as e:
            log(f"WARN:Playwright WoS retractados error: {e}")
        finally:
            browser.close()

    return retracted_count


def get_wos_journal_data(issn, wos_api_key=None):
    """
    Obtiene JIF, cuartil, categorías JCR y colección WoS para una revista.

    Estrategia con dos niveles:
      1) WoS Journals API (Clarivate) — si se provee wos_api_key
         GET https://api.clarivate.com/apis/wos-journals/v1/journals?q={issn}&jcrYear={año}
      2) Fallback: scraping de wos-journal.info (sin autenticación)
         GET https://wos-journal.info/journalid/{issn}

    Devuelve dict:
      {
        "found":       True/False,
        "source":      "wos_api" | "wos_journal_info" | None,
        "jif":         "5.123",
        "jif_year":    "2024",
        "collections": ["SCIE"],           # SCIE / SSCI / ESCI / AHCI
        "categories":  [
            {"name": "Pharmaceutics", "quartile": "Q1", "rank": "12/165"},
            ...
        ],
      }
    """
    if not issn:
        return {"found": False, "source": None}

    clean = issn.replace("-", "").strip()
    issn_fmt = f"{clean[:4]}-{clean[4:]}" if len(clean) == 8 else issn.strip()
    jcr_year = str(date.today().year - 1)  # JCR del año anterior (el más reciente disponible)

    # ── Nivel 1: WoS Journals API oficial ──────────────────────────────────
    if wos_api_key:
        try:
            url = "https://api.clarivate.com/apis/wos-journals/v1/journals"
            headers = {"X-ApiKey": wos_api_key, "Accept": "application/json"}
            params  = {"q": issn_fmt, "jcrYear": jcr_year, "limit": 5}
            resp = requests.get(url, headers=headers, params=params, timeout=15)
            if resp.status_code == 200:
                data    = resp.json()
                hits    = data.get("hits", [])
                journal = next(
                    (j for j in hits
                     if issn_fmt in (j.get("issn",""), j.get("eissn",""))),
                    hits[0] if hits else None
                )
                if journal:
                    metrics = journal.get("metrics", {})
                    ranks   = journal.get("ranks",   {})
                    # Colecciones
                    editions = [e.get("edition","") for e in journal.get("coverageHistory", [])]
                    if not editions:
                        editions = journal.get("editions", [])

                    # Categorías + cuartiles
                    categories = []
                    for cat in ranks.get("jif", []):
                        categories.append({
                            "name":     cat.get("category",""),
                            "quartile": cat.get("quartile",""),
                            "rank":     f"{cat.get('rank','')}/{cat.get('total','')}",
                        })

                    return {
                        "found":       True,
                        "source":      "wos_api",
                        "jif":         str(metrics.get("jif", {}).get("current", "")),
                        "jif_year":    jcr_year,
                        "collections": list(set(editions)) or ["Ver JCR"],
                        "categories":  categories,
                    }
        except Exception as e:
            log(f"WARN:WoS API: {e}")

    # ── Nivel 2: wos-journal.info (scraping sin autenticación) ─────────────
    try:
        url  = f"https://wos-journal.info/journalid/{issn_fmt}"
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=20)
        if resp.status_code != 200:
            # Intentar sin guión
            url  = f"https://wos-journal.info/journalid/{clean}"
            resp = requests.get(url, headers=HTTP_HEADERS, timeout=20)

        if resp.status_code == 200:
            soup = BeautifulSoup(resp.text, "html.parser")

            # JIF
            jif_val  = ""
            jif_year = ""
            for tag in soup.select("td, th, span, div"):
                txt = tag.get_text(strip=True)
                if "Journal Impact Factor" in txt or "JIF" in txt:
                    # El valor suele estar en el siguiente elemento
                    nxt = tag.find_next(string=re.compile(r"^\d+\.\d+$"))
                    if nxt:
                        jif_val = nxt.strip()
                        break

            # Buscar JIF en tablas de métricas
            if not jif_val:
                for row in soup.select("table tr"):
                    cells = [c.get_text(strip=True) for c in row.select("td,th")]
                    for i, c in enumerate(cells):
                        if re.match(r"^\d+\.\d+$", c) and i > 0:
                            label = cells[i-1].upper()
                            if "JIF" in label or "IMPACT" in label:
                                jif_val = c; break

            # Colecciones WoS
            collections = []
            for badge in soup.select(".badge, .label, .tag, span"):
                t = badge.get_text(strip=True).upper()
                if t in ("SCIE","SSCI","ESCI","AHCI"):
                    if t not in collections:
                        collections.append(t)
            # Fallback: buscar en texto
            if not collections:
                page_text = soup.get_text()
                for col in ("SCIE","SSCI","ESCI","AHCI"):
                    if col in page_text:
                        collections.append(col)

            # Categorías y cuartiles — buscar tabla o lista con Q1/Q2/Q3/Q4
            categories = []
            for row in soup.select("table tr"):
                cells = [c.get_text(strip=True) for c in row.select("td,th")]
                for i, c in enumerate(cells):
                    if re.match(r"^Q[1-4]$", c):
                        cat_name = cells[i-1] if i > 0 else ""
                        rank_str = cells[i+1] if i+1 < len(cells) else ""
                        if cat_name and cat_name not in ("", "Category","Categoría"):
                            categories.append({
                                "name":     cat_name,
                                "quartile": c,
                                "rank":     rank_str,
                            })

            if jif_val or collections or categories:
                return {
                    "found":       True,
                    "source":      "wos_journal_info",
                    "jif":         jif_val,
                    "jif_year":    jif_year or jcr_year,
                    "collections": collections or ["Ver JCR"],
                    "categories":  categories,
                }
    except Exception as e:
        log(f"WARN:wos-journal.info: {e}")

    # ── Nivel 3: Playwright → JCR directo (más confiable, requiere acceso) ──
    pw_result = scrape_jcr_playwright(issn_fmt)
    if pw_result and pw_result.get("found"):
        return pw_result

    return {"found": False, "source": None}


PREDATORY_SHEET_ID = "1Qa1lAlSbl7iiKddYINNsDB4wxI7uUA4IVseeLnCc5U4"
HTTP_HEADERS = {"User-Agent": "Mozilla/5.0 (compatible; research-bot/1.0)"}

def _normalize(text):
    if not text: return ""
    text = str(text).lower().strip()
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()

def _fetch_html(url):
    try:
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=20)
        if resp.status_code == 200:
            return BeautifulSoup(resp.text, "html.parser")
    except: pass
    return None

def _parse_bealls(url):
    result = {}
    soup = _fetch_html(url)
    if not soup: return result
    for a in soup.select("ul li a[href]"):
        href = a.get("href", "")
        text = a.get_text(strip=True)
        if text and len(text) > 3 and href.startswith("http") and "beallslist.net" not in href:
            clean = re.sub(r'\s*\(.*?\)', '', text).strip()
            if clean:
                result[_normalize(clean)] = clean
    return result

def _parse_predatory_sheet():
    result = {}
    url = f"https://docs.google.com/spreadsheets/d/{PREDATORY_SHEET_ID}/export?format=csv"
    try:
        resp = requests.get(url, headers=HTTP_HEADERS, timeout=30)
        if resp.status_code != 200: return result
        df = pd.read_csv(StringIO(resp.text), dtype=str, header=None)
        for val in df.values.flatten():
            if val is None or pd.isna(val): continue
            val = str(val).strip()
            if not val or re.fullmatch(r'\d+', val) or len(val) < 4: continue
            result[_normalize(val)] = val
    except: pass
    return result

def check_predatory(journal_name, publisher):
    publishers_list = _parse_bealls("https://beallslist.net/")
    time.sleep(0.5)
    standalone_list = _parse_bealls("https://beallslist.net/standalone-journals/")
    time.sleep(0.5)
    predatory_list  = _parse_predatory_sheet()

    found = []
    if journal_name:
        if predatory_list.get(_normalize(journal_name)):
            found.append("Lista 1 (predatoryjournals.org)")
        if standalone_list.get(_normalize(journal_name)):
            found.append("Lista 2 - Beall's (journal)")
    if publisher:
        m = publishers_list.get(_normalize(publisher))
        if m: found.append(f"Lista 2 - Beall's (publisher: {m})")
    return found

def generate_year_chart(pubs_by_year):
    if not pubs_by_year: return None
    years  = sorted(pubs_by_year.keys())
    counts = [pubs_by_year[y] for y in years]
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.plot(years, counts, marker="o", color="#2E75B6", linewidth=2.5, markersize=7)
    for x, y in zip(years, counts):
        ax.annotate(f"{y:,}", (x, y), textcoords="offset points",
                    xytext=(0, 10), ha="center", fontsize=9)
    ax.set_xlabel("Año", fontsize=10)
    ax.set_ylabel("Documentos", fontsize=10)
    ax.set_xticks(years)
    ax.set_xticklabels([str(y) for y in years], rotation=0, fontsize=9)
    ax.yaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def generate_country_chart(pubs_by_country: list, top_n=15) -> BytesIO:
    """
    Gráfico de barras horizontales con los top países.
    pubs_by_country: [{"country": ..., "count": ...}, ...]
    """
    if not pubs_by_country:
        return None

    data = pubs_by_country[:top_n]
    data = sorted(data, key=lambda x: x["count"])  # menor arriba

    countries = [d["country"] for d in data]
    counts    = [d["count"]   for d in data]

    # Paleta de azules degradados
    n = len(countries)
    colors = [plt.cm.Blues(0.4 + 0.55 * (i / max(n - 1, 1))) for i in range(n)]

    fig, ax = plt.subplots(figsize=(9, max(3.5, n * 0.38)))
    bars = ax.barh(countries, counts, color=colors, edgecolor="white", height=0.65)

    # Etiquetas al extremo derecho de cada barra
    for bar, val in zip(bars, counts):
        ax.text(bar.get_width() + max(counts) * 0.01, bar.get_y() + bar.get_height() / 2,
                f"{val:,}", va="center", ha="left", fontsize=9)

    ax.set_xlabel("Documentos", fontsize=10)
    ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    ax.set_xlim(0, max(counts) * 1.15)
    ax.grid(axis="x", linestyle="--", alpha=0.35)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


# ─── WORD helpers ──────────────────────────────────────────────────────────────
YELLOW = RGBColor(0xFF, 0xFF, 0x00)
GRAY   = RGBColor(0xBF, 0xBF, 0xBF)

def highlight_run(run):
    rPr = run._r.get_or_add_rPr()
    hl  = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    rPr.append(hl)

def add_caption(doc, text):
    """Agrega nota: solo 'Nota:' en cursiva, el resto en texto normal."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    if text.startswith("Nota:"):
        r_nota = p.add_run("Nota:")
        r_nota.italic = True
        r_nota.font.size = Pt(9)
        r_nota.font.name = "Times New Roman"
        r_rest = p.add_run(text[len("Nota:"):])
        r_rest.font.size = Pt(9)
        r_rest.font.name = "Times New Roman"
    else:
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

def add_hyperlink(paragraph, url, text, font_name="Times New Roman", font_size_pt=11):
    """Agrega un hipervínculo con color RGB(0,0,255) y sin subrayado."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hl_run = OxmlElement("w:r")
    rPr_hl = OxmlElement("w:rPr")
    # Font
    rFonts_hl = OxmlElement("w:rFonts")
    rFonts_hl.set(qn("w:ascii"), font_name)
    rFonts_hl.set(qn("w:hAnsi"), font_name)
    rFonts_hl.set(qn("w:eastAsia"), font_name)
    rFonts_hl.set(qn("w:cs"), font_name)
    rPr_hl.append(rFonts_hl)
    # Size
    sz_hl = OxmlElement("w:sz")
    sz_hl.set(qn("w:val"), str(font_size_pt * 2))  # half-points
    rPr_hl.append(sz_hl)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size_pt * 2))
    rPr_hl.append(szCs)
    # Color RGB(0,0,255)
    color_hl = OxmlElement("w:color")
    color_hl.set(qn("w:val"), "0000FF")
    rPr_hl.append(color_hl)
    # Sin subrayado
    u_hl = OxmlElement("w:u")
    u_hl.set(qn("w:val"), "none")
    rPr_hl.append(u_hl)
    hl_run.append(rPr_hl)
    hl_text = OxmlElement("w:t")
    hl_text.set(qn("xml:space"), "preserve")
    hl_text.text = text
    hl_run.append(hl_text)
    hyperlink.append(hl_run)
    paragraph._p.append(hyperlink)

def add_image_border(doc):
    """Agrega borde de 1pt color #999999 a la última imagen insertada."""
    last_paragraph = doc.paragraphs[-1]
    inline_elements = last_paragraph._p.findall(
        './/' + qn('wp:inline')
    )
    if not inline_elements:
        # Try anchor
        inline_elements = last_paragraph._p.findall(
            './/' + qn('wp:anchor')
        )
    for inline in inline_elements:
        spPr = inline.find('.//' + qn('pic:spPr'))
        if spPr is None:
            continue
        # Remove existing line if present
        existing_ln = spPr.find(qn('a:ln'))
        if existing_ln is not None:
            spPr.remove(existing_ln)
        # Add 1pt solid line with #999999
        ln = OxmlElement('a:ln')
        ln.set('w', '12700')  # 1pt = 12700 EMU
        solidFill = OxmlElement('a:solidFill')
        srgbClr = OxmlElement('a:srgbClr')
        srgbClr.set('val', '999999')
        solidFill.append(srgbClr)
        ln.append(solidFill)
        spPr.append(ln)

def generate_word_report(meta, pubs_by_year, total_docs, retracted_scopus,
                         predatory_hits, year_chart_buf, country_chart_buf,
                         pubs_by_country, output_file, nro_informe, institucion,
                         retracted_wos=None):
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for h_level, pt_size, indent_cm, before, after in [
        (1, 11, 0.0, 12, 6),
        (2, 11, 0.5, 6, 4),
        (3, 11, 1.0, 3, 2)
    ]:
        try:
            style = doc.styles[f"Heading {h_level}"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(pt_size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            style.paragraph_format.left_indent = Cm(indent_cm)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Force all rFonts attributes at XML level to override Word theme
            rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
            if rFonts is None:
                rPr = style.element.get_or_add_rPr()
                rFonts_el = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts_el)
                rFonts = rFonts_el
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
        except: pass

    def add_heading_tnr(text, level=1):
        """Add a heading with Times New Roman explicitly set on the run."""
        h = doc.add_paragraph(style=f"Heading {level}")
        r = h.add_run(text)
        r.font.name = "Times New Roman"
        # Also set at XML level on the run itself
        rPr = r._r.get_or_add_rPr()
        rFonts_el = OxmlElement('w:rFonts')
        rFonts_el.set(qn('w:ascii'), 'Times New Roman')
        rFonts_el.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts_el.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts_el.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts_el)
        return h

    # ── Encabezado ──────────────────────────────────────────────────────────
    for line in institucion:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        r.font.name = "Times New Roman"; r.font.size = Pt(11); r.font.bold = True
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Informe N° IDIC-IQInteg-R-")
    r.font.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(nro_informe)
    r2.font.bold = True; r2.font.size = Pt(11)
    if nro_informe.startswith("["):
        highlight_run(r2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Integridad científica de la revista ")
    r.font.bold = True; r.font.size = Pt(11)
    rj = p.add_run(meta['journal'])
    rj.font.bold = True; rj.font.size = Pt(11); rj.font.italic = True

    today = date.today()
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    fecha_str = f"Lima, {today.day:02d} de {meses[today.month-1]} de {today.year}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(fecha_str); r.font.size = Pt(11); r.font.name = "Times New Roman"

    # ── Artículo de referencia (si se ingresó por Scopus ID) ─────────────────
    if meta.get("scopus_id"):
        p_hdr = doc.add_paragraph(style="Normal")
        p_hdr.paragraph_format.space_after = Pt(6)
        r_hdr = p_hdr.add_run("Artículo de referencia analizado")
        r_hdr.bold = True; r_hdr.font.size = Pt(11)

        for lbl, val in [
            ("Scopus ID",  meta.get("scopus_id", "")),
            ("Título",     meta.get("article_title", "")),
            ("Autor(es)",  meta.get("article_author", "")),
            ("DOI",        meta.get("article_doi", "")),
            ("Fecha",      meta.get("article_date", "")),
        ]:
            if val:
                p2 = doc.add_paragraph(style="Normal")
                p2.paragraph_format.left_indent = Cm(0.5)
                p2.paragraph_format.space_after = Pt(2)
                r_l = p2.add_run(f"{lbl}: "); r_l.bold = True
                p2.add_run(str(val))

    # ── 1. Datos de la revista ──────────────────────────────────────────────
    add_heading_tnr("1. Datos de la revista", level=1)

    # 1.1 ISSN
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.1 ISSN: "); r.bold = True
    issn_val = meta.get("issn_print", "") or "[no disponible]"
    p.add_run(f"{issn_val};  E-ISSN: {meta['eissn']}")

    # 1.2
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.2 Sitio web: "); r.bold = True
    homepage_val = meta.get("homepage") or "[no disponible]"
    if homepage_val.startswith("http"):
        add_hyperlink(p, homepage_val, homepage_val)
    else:
        p.add_run(homepage_val)

    # 1.3
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.3 Editorial: "); r.bold = True
    p.add_run(meta["publisher"])

    # 1.4 APC — from editable or placeholder
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.4 APC: "); r.bold = True
    apc_val = meta.get("apc", "").strip()
    if apc_val and not apc_val.startswith("["):
        p.add_run(apc_val)
    else:
        rv = p.add_run("[COMPLETAR: monto y moneda, o 'Sin APC']")
        highlight_run(rv)

    # 1.5 Tiempo — from editable or placeholder
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.5 Tiempo de publicación de artículos: "); r.bold = True
    pt_val = meta.get("pub_time", "").strip()
    if pt_val and not pt_val.startswith("["):
        p.add_run(pt_val)
    else:
        rv = p.add_run("[COMPLETAR: días promedio o 'No precisa el sitio web']")
        highlight_run(rv)

    # 1.6 CiteScore
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"1.6 Cuartil CiteScore (Scopus {meta['citescore_year']}):"); r.bold = True

    if meta["quartiles"]:
        for area, subareas in meta["quartiles"].items():
            p2 = doc.add_paragraph(style="Normal")
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.paragraph_format.space_after  = Pt(2)
            p2_run = p2.add_run(f"De acuerdo con el área temática ")
            p2_area = p2.add_run(area)
            p2_area.italic = True
            p2.add_run(", la revista se clasifica en las siguientes categorías y cuartiles:")
            for item in subareas:
                p_cs = doc.add_paragraph(style="Normal")
                p_cs.paragraph_format.left_indent = Cm(1)
                p_cs.add_run(f"Se ubica en el cuartil {item['cuartil'].replace('Q','')} ")
                r_q = p_cs.add_run(f"({item['cuartil']})")
                r_q.bold = True
                p_cs.add_run(" en la categoría ")
                r_s = p_cs.add_run(item['subarea'])
                r_s.italic = True
                p_cs.add_run(".")
    else:
        p2 = doc.add_paragraph(style="Normal")
        p2.paragraph_format.left_indent = Cm(0.5)
        rv = p2.add_run("[No se encontraron datos de cuartiles CiteScore]")
        highlight_run(rv)

    # 1.7 WoS — Colecciones y categorias JCR (via Starter API)
    wos_collections = meta.get("wos_collections", [])
    wos_categories  = meta.get("wos_categories", [])
    jcr_year_label  = str(date.today().year - 2)
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"1.7 Cuartil WoS (JIF {jcr_year_label}):"); r.bold = True

    if wos_categories or wos_collections:
        # Categorias JCR primero — con espacio para cuartiles
        if wos_categories:
            for cat_name in wos_categories:
                wos_q_map = meta.get("wos_quartiles", {})
                q_val = wos_q_map.get(cat_name, "").strip() if isinstance(wos_q_map, dict) else ""
                p_cat = doc.add_paragraph(style="Normal")
                p_cat.paragraph_format.left_indent = Cm(0.5)
                p_cat.paragraph_format.space_after = Pt(2)
                p_cat.add_run("Se ubica en el cuartil ")
                if q_val and not q_val.startswith("["):
                    r_q = p_cat.add_run(q_val)
                    r_q.bold = True
                else:
                    rv_q = p_cat.add_run("[Q?]")
                    highlight_run(rv_q)
                p_cat.add_run(" en la categoria ")
                r_cat = p_cat.add_run(cat_name)
                r_cat.italic = True
                p_cat.add_run(".")
        else:
            p_nc = doc.add_paragraph(style="Normal")
            p_nc.paragraph_format.left_indent = Cm(0.5)
            rv_nc = p_nc.add_run("[COMPLETAR: cuartil y categoria JCR]")
            highlight_run(rv_nc)

        # Colecciones al final
        wosm = meta.get("wos_collections_manual", "").strip()
        if wos_collections or wosm:
            all_cols = list(wos_collections)
            if wosm:
                # split by comma, clean, add
                for c in wosm.split(','):
                    cc = c.strip()
                    if cc and cc not in all_cols: all_cols.append(cc)
            
            for col_name in all_cols:
                p_col = doc.add_paragraph(style="Normal")
                p_col.paragraph_format.left_indent = Cm(0.5)
                p_col.paragraph_format.space_after = Pt(2)
                p_col.add_run("Forma parte de la colección ")
                r_col = p_col.add_run(col_name)
                r_col.italic = True
                p_col.add_run(".")
        else:
            p_col = doc.add_paragraph(style="Normal")
            p_col.paragraph_format.left_indent = Cm(0.5)
            rv_col = p_col.add_run("[COMPLETAR: coleccion WoS, ej: SCIE / SSCI / ESCI]")
            highlight_run(rv_col)
    else:
        p2 = doc.add_paragraph(style="Normal"); p2.paragraph_format.left_indent = Cm(0.5)
        rv = p2.add_run("[COMPLETAR: cuartil y categoria JCR, ej: Q2 en Engineering, Civil]")
        highlight_run(rv)
        p3 = doc.add_paragraph(style="Normal"); p3.paragraph_format.left_indent = Cm(0.5)
        rv2 = p3.add_run("[COMPLETAR: coleccion WoS, ej: SCIE / SSCI / ESCI]")
        highlight_run(rv2)

    # 1.8 Vigencia Scopus
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("1.8 Vigencia en Scopus: "); r.bold = True
    p.add_run(
        f"Revista vigente en Scopus de {meta['coverage_start']} a {meta['coverage_end']}."
    )
    if meta.get("scopus_link"):
        p2 = doc.add_paragraph(style="Normal")
        p2.paragraph_format.left_indent = Cm(0.5)
        add_hyperlink(p2, meta["scopus_link"], meta["scopus_link"])

    doc.add_paragraph()

    # ── 2. Producción científica ─────────────────────────────────────────────
    add_heading_tnr("2. Producción científica de la revista", level=1)

    add_heading_tnr("2.1 Publicaciones por año", level=2)

    if total_docs > 0:
        p = doc.add_paragraph(style="Normal"); p.paragraph_format.space_after = Pt(6)
        p.add_run(
            f"La revista {meta['journal']} registra, a la fecha, "
            f"{total_docs:,} documentos indexados en Scopus."
        )

    if pubs_by_year:
        years_sorted = sorted(pubs_by_year.keys())
        y_start = years_sorted[0]; y_end = years_sorted[-1]
        c_start = pubs_by_year[y_start]
        c_max   = max(pubs_by_year.values())
        y_max   = max(pubs_by_year, key=pubs_by_year.get)
        p = doc.add_paragraph(style="Normal"); p.paragraph_format.space_after = Pt(6)
        p.add_run(
            f"Durante el período {y_start}–{y_end}, el volumen inició en {c_start:,} "
            f"documentos en {y_start} y alcanzó su punto más alto en {y_max} con {c_max:,} documentos."
        )

    if year_chart_buf:
        p_fig = doc.add_paragraph(style="Normal")
        p_fig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_fig.paragraph_format.space_before = Pt(6)
        p_fig.add_run("Figura 1").bold = True
        p_fig.add_run("\n")
        p_fig.add_run("Publicaciones por año de la revista").italic = True
        doc.add_picture(year_chart_buf, width=Cm(14))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_border(doc)
        add_caption(doc, "Nota: Elaborado a partir de datos extraídos de Scopus.")

    doc.add_paragraph()

    # 2.2 Por país — AUTOMÁTICO
    add_heading_tnr("2.2 Documentos por país", level=2)

    if pubs_by_country:
        top3 = pubs_by_country[:3]
        top_names = ", ".join(d["country"] for d in top3)
        total_countries = len(pubs_by_country)
        total_country_docs = sum(d["count"] for d in pubs_by_country)
        p = doc.add_paragraph(style="Normal"); p.paragraph_format.space_after = Pt(6)
        p.add_run(
            f"La revista registra contribuciones de {total_countries} países. "
            f"Los países con mayor volumen de publicaciones son {top_names}, "
            f"concentrando la mayoría de los {total_country_docs:,} documentos analizados."
        )
    else:
        p = doc.add_paragraph(style="Normal"); p.paragraph_format.space_after = Pt(6)
        rv = p.add_run("[No se obtuvieron datos de países desde Scopus]")
        highlight_run(rv)

    if country_chart_buf:
        p_fig2 = doc.add_paragraph(style="Normal")
        p_fig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_fig2.paragraph_format.space_before = Pt(6)
        p_fig2.add_run("Figura 2").bold = True
        p_fig2.add_run("\n")
        p_fig2.add_run("Publicaciones de la revista por país").italic = True
        doc.add_picture(country_chart_buf, width=Cm(14))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_border(doc)
        add_caption(doc, "Nota: Elaborado a partir de datos extraídos de Scopus.")
    else:
        p_ph = doc.add_paragraph(style="Normal"); p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = p_ph.add_run("[No se pudo generar gráfico de países]")
        rv.font.color.rgb = GRAY; rv.italic = True

    doc.add_paragraph()

    # 2.3 Retractados
    add_heading_tnr("2.3 Artículos retractados por la revista", level=2)

    # Scopus
    if retracted_scopus == 0:
        rtxt_scopus = "No se identificaron artículos retractados en Scopus."
    elif retracted_scopus == 1:
        rtxt_scopus = "Se identificó un (1) artículo retractado en Scopus."
    else:
        rtxt_scopus = f"Se identificaron {retracted_scopus} artículos retractados en Scopus."

    p = doc.add_paragraph(style="Normal"); p.add_run(rtxt_scopus)

    # WoS — automático si Playwright pudo obtenerlo, manual si no
    p2 = doc.add_paragraph(style="Normal")
    if retracted_wos is not None:
        if retracted_wos == 0:
            rtxt_wos = "No se identificaron artículos retractados en Web of Science."
        elif retracted_wos == 1:
            rtxt_wos = "Se identificó un (1) artículo retractado en Web of Science."
        else:
            rtxt_wos = f"Se identificaron {retracted_wos} artículos retractados en Web of Science."
        p2.add_run(rtxt_wos)
    else:
        wos_manual = meta.get("retracted_wos_manual", "").strip()
        if wos_manual:
            if wos_manual.isdigit():
                n = int(wos_manual)
                if n == 0:
                    p2.add_run("No se identificaron artículos retractados en Web of Science.")
                else:
                    p2.add_run(f"Se identificaron {n} artículos retractados en Web of Science.")
            else:
                p2.add_run(wos_manual)
        else:
            rv = p2.add_run("[COMPLETAR: artículos retractados en Web of Science]")
            highlight_run(rv)

    doc.add_paragraph()

    # ── 3. Listas depredadoras ───────────────────────────────────────────────
    add_heading_tnr("3. Verificación de listas de revistas predatorias", level=1)

    lista1_url = "https://www.predatoryjournals.org/the-list/journals"
    lista2_url = "https://beallslist.net/"

    if not predatory_hits:
        p = doc.add_paragraph(style="Normal")
        p.add_run("No figura en la lista de Predatory journals ni en Beall's List.")
        bp1 = doc.add_paragraph(style="List Bullet")
        add_hyperlink(bp1, lista1_url, lista1_url)
        bp2 = doc.add_paragraph(style="List Bullet")
        add_hyperlink(bp2, lista2_url, lista2_url)
    else:
        p = doc.add_paragraph(style="Normal")
        r = p.add_run("⚠ La revista figura en las siguientes listas de revistas predatorias:")
        r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        for hit in predatory_hits:
            bp = doc.add_paragraph(style="List Bullet")
            rh = bp.add_run(hit); rh.bold = True
        p2 = doc.add_paragraph(style="Normal"); p2.add_run("Fuentes verificadas:")
        bp1 = doc.add_paragraph(style="List Bullet")
        add_hyperlink(bp1, lista1_url, lista1_url)
        bp2 = doc.add_paragraph(style="List Bullet")
        add_hyperlink(bp2, lista2_url, lista2_url)

    doc.add_paragraph()

    # ── 4. Conclusiones ─────────────────────────────────────────────────────
    add_heading_tnr("4. Conclusiones", level=1)

    # 4.1
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run("4.1 ").bold = True
    p.add_run(
        f"Es una revista indexada en la base de datos Scopus "
        f"(vigente de {meta['coverage_start']} a {meta['coverage_end']})."
    )

    if meta["quartiles"]:
        all_q = [item["cuartil"] for sl in meta["quartiles"].values() for item in sl]
        best_q = sorted(set(all_q))[0] if all_q else "N/A"
        c2 = (f"Presenta métricas CiteScore con clasificación en {best_q} "
              f"en sus categorías de mayor desempeño.")
    else:
        c2 = meta.get("concl_metrics", "").strip()
        if not c2:
            c2 = "[COMPLETAR: describir desempeño en métricas]"

    # 4.2
    p2 = doc.add_paragraph(style="Normal")
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run("4.2 ").bold = True
    rv2 = p2.add_run(c2)
    if c2.startswith("["): highlight_run(rv2)

    c3 = ("No presenta indicios de malas prácticas (no figura en listas de revistas predatorias)."
          if not predatory_hits
          else "⚠ Figura en listas de revistas predatorias. Se recomienda evaluar con cautela.")
    # 4.3
    p3 = doc.add_paragraph(style="Normal")
    p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run("4.3 ").bold = True
    rv3 = p3.add_run(c3)
    if predatory_hits: rv3.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # ── 5. Recomendaciones ──────────────────────────────────────────────────
    add_heading_tnr("5. Recomendaciones", level=1)

    # 5.1
    PROTOCOLO_URL = "https://www.ulima.edu.pe/sites/default/files/centers/files/protocolo_idoneidad_publicaciones_idic_v25jul_1.pdf"
    p51 = doc.add_paragraph(style="Normal")
    p51.paragraph_format.left_indent = Cm(0.5)
    p51.paragraph_format.space_before = Pt(4)
    p51.paragraph_format.space_after = Pt(2)
    p51.add_run("5.1 ").bold = True
    p51.add_run("Para elecciones futuras de una revista científica, se recomienda leer el ")

    # Add hyperlink for the Protocolo using the helper
    add_hyperlink(p51, PROTOCOLO_URL, "Protocolo para la evaluación de la idoneidad de publicaciones científicas")

    p51.add_run(
        ". Esta revisión permitirá comprender mejor los criterios de calidad que la "
        "institución considera importantes al evaluar las publicaciones científicas."
    )

    # 5.2
    p52 = doc.add_paragraph(style="Normal")
    p52.paragraph_format.left_indent = Cm(0.5)
    p52.paragraph_format.space_before = Pt(4)
    p52.paragraph_format.space_after = Pt(2)
    p52.add_run("5.2 ").bold = True
    p52.add_run(
        "Es importante reconocer que el comportamiento editorial, la calidad y la "
        "reputación de una revista pueden cambiar con el tiempo. Por lo tanto, se "
        "aconseja verificar sus indexaciones periódicamente."
    )

    doc.add_paragraph()

    # Pie de página
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp.add_run(
        "El presente informe tiene una vigencia de tres (03) meses "
        "a partir de su fecha de emisión."
    )
    rf.font.size = Pt(9); rf.italic = True
    rf.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.save(output_file)


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    params = json.loads(sys.stdin.read())
    mode = params.get("mode", "full")  # "collect", "generate", or "full"

    api_key = params["api_key"]

    # ══════════════════════════════════════════════════════════════════════════
    #  MODE: COLLECT — Gather all data, stream progress, output JSON at end
    # ══════════════════════════════════════════════════════════════════════════
    if mode == "collect":
        end_year   = date.today().year
        start_year = end_year - 8

        scopus_id = params.get("scopus_id", "").strip()
        eissn     = params.get("eissn", "").strip()
        article_info = None

        if scopus_id:
            log("STEP:article")
            log(f"INFO:Resolviendo Scopus ID: {scopus_id}")
            article_info = get_article_and_journal_from_scopus_id(api_key, scopus_id)
            log(f"INFO:Articulo: {article_info['article_title'][:80]}")
            log(f"INFO:Revista detectada: {article_info['journal_name']}")
            eissn = article_info["issn_for_lookup"]
            if not eissn:
                raise RuntimeError("No se pudo obtener ISSN de la revista del articulo")
            log(f"INFO:ISSN para busqueda de revista: {eissn}")

        if not eissn:
            raise RuntimeError("Debes proporcionar 'scopus_id' o 'eissn'")

        log("STEP:metadata")
        log("INFO:Consultando metadatos de la revista...")
        meta = get_journal_metadata(api_key, eissn)
        if article_info:
            meta["article_title"]  = article_info["article_title"]
            meta["article_doi"]    = article_info["article_doi"]
            meta["article_date"]   = article_info["article_date"]
            meta["article_author"] = article_info["article_author"]
            meta["scopus_id"]      = scopus_id
            if not meta.get("source_id") and article_info.get("source_id"):
                meta["source_id"] = article_info["source_id"]
        log(f"INFO:Revista: {meta['journal']}")
        # Send metadata to frontend for preview
        meta_out = dict(meta)
        meta_out["_key"] = "meta"
        log(f"DATA:{json.dumps(meta_out, ensure_ascii=False)}")

        log("STEP:total_docs")
        log("INFO:Contando documentos totales...")
        total_docs = get_total_documents(api_key, meta["source_id"])
        log(f"INFO:{total_docs:,} documentos totales")
        log(f'DATA:{json.dumps({"_key": "total_docs", "value": total_docs})}')

        log("STEP:pubs_by_year")
        log("INFO:Consultando publicaciones por ano...")
        pubs_by_year = get_publications_by_year(api_key, meta["source_id"], start_year, end_year)
        log(f"INFO:{len(pubs_by_year)} anos con datos")
        log(f'DATA:{json.dumps({"_key": "pubs_by_year", "value": {str(k): v for k, v in pubs_by_year.items()}})}')

        log("STEP:retracted")
        log("INFO:Buscando articulos retractados en Scopus...")
        retracted = get_retracted_count(api_key, meta["source_id"])
        log(f"INFO:{retracted} articulos retractados en Scopus")
        log(f'DATA:{json.dumps({"_key": "retracted_scopus", "value": retracted})}')

        log("STEP:wos_collections")
        log("INFO:Consultando colecciones WoS via Starter API...")
        wos_api_key  = params.get("wos_api_key", "").strip()
        _init_wos_api(wos_api_key)
        eissn_val = meta.get("eissn") or ""
        issn_print_val = meta.get("issn_print") or ""
        wos_collections = get_wos_collections(eissn_val, issn_print_val)
        if wos_collections:
            log(f"INFO:{len(wos_collections)} colecciones WoS encontradas")
        else:
            log("WARN:No se encontraron colecciones WoS (sin API key o revista no indexada)")
        log(f'DATA:{json.dumps({"_key": "wos_collections", "value": wos_collections}, ensure_ascii=False)}')

        log("STEP:wos_categories")
        log("INFO:Consultando categorias JCR via Starter API...")
        wos_categories = get_wos_categories(eissn_val, issn_print_val)
        if wos_categories:
            log(f"INFO:{len(wos_categories)} categorias JCR encontradas")
        else:
            log("WARN:No se encontraron categorias JCR")
        log(f'DATA:{json.dumps({"_key": "wos_categories", "value": wos_categories}, ensure_ascii=False)}')

        log("STEP:wos_retracted")
        log("INFO:Contando retractados en WoS via Starter API...")
        retracted_wos_val = get_wos_retracted_count(eissn_val, issn_print_val)
        if retracted_wos_val >= 0:
            log(f"INFO:Retractados WoS: {retracted_wos_val}")
        else:
            log("WARN:No se pudo obtener retractados WoS")
        log(f'DATA:{json.dumps({"_key": "retracted_wos", "value": retracted_wos_val})}')

        log("STEP:predatory")
        log("INFO:Verificando listas de revistas predatorias...")
        hits = check_predatory(meta["journal"], meta["publisher"])
        if hits:
            log(f"WARN:DEPREDADORA detectada: {hits}")
        else:
            log("INFO:No figura en listas depredadoras")
        log(f'DATA:{json.dumps({"_key": "predatory", "value": hits}, ensure_ascii=False)}')

        log("STEP:countries")
        log("INFO:Consultando publicaciones por pais...")
        pubs_by_country = get_publications_by_country(api_key, meta.get("source_id"), top_n=15)
        log(f"INFO:{len(pubs_by_country)} paises con datos")
        log(f'DATA:{json.dumps({"_key": "pubs_by_country", "value": pubs_by_country}, ensure_ascii=False)}')

        # Final collected data bundle
        collected = {
            "meta": meta,
            "total_docs": total_docs,
            "pubs_by_year": {str(k): v for k, v in pubs_by_year.items()},
            "retracted_scopus": retracted,
            "retracted_wos": retracted_wos_val,
            "wos_collections": wos_collections,
            "wos_categories": wos_categories,
            "predatory": hits,
            "pubs_by_country": pubs_by_country,
        }
        log(f"COLLECT_DONE:{json.dumps(collected, ensure_ascii=False)}")

    # ══════════════════════════════════════════════════════════════════════════
    #  MODE: GENERATE — Take reviewed data and produce .docx
    # ══════════════════════════════════════════════════════════════════════════
    elif mode == "generate":
        rd = params.get("report_data", {})
        meta           = rd.get("meta")
        if not meta:
            # Fallback for old flat structure if absolutely necessary, 
            # but usually meta is required.
            raise ValueError("El objeto 'report_data' debe contener una clave 'meta'.")
        total_docs     = rd.get("total_docs", 0)
        pubs_by_year   = {int(k): v for k, v in rd.get("pubs_by_year", {}).items()}
        retracted      = rd.get("retracted_scopus", 0)
        retracted_wos  = rd.get("retracted_wos")
        if retracted_wos == -1:
            retracted_wos = None
        meta["wos_collections"] = rd.get("wos_collections", [])
        meta["wos_categories"]  = rd.get("wos_categories", [])
        predatory_hits     = rd.get("predatory", [])
        pubs_by_country    = rd.get("pubs_by_country", [])

        # Inject editable fields from user
        meta["apc"]              = rd.get("apc", "")
        meta["pub_time"]         = rd.get("pub_time", "")
        meta["retracted_wos_manual"] = rd.get("retracted_wos_manual", "")
        meta["wos_quartiles"]    = rd.get("wos_quartiles", {})
        meta["wos_collections_manual"] = rd.get("wos_collections_manual", "")
        meta["concl_metrics"]    = rd.get("concl_metrics", "")
        # Override homepage if user edited it
        homepage_edit = rd.get("homepage", "").strip()
        if homepage_edit:
            meta["homepage"] = homepage_edit

        nro_informe = params.get("nro_informe", "[COMPLETAR]")
        institucion = params.get("institucion", [
            "Área de Inteligencia e Integridad",
            "Instituto de Investigación Científica",
            "Universidad de Lima",
        ])
        output_file = params.get("output_file", "informe.docx")

        log("STEP:chart")
        chart_buf = generate_year_chart(pubs_by_year)

        log("STEP:country_chart")
        country_chart_buf = generate_country_chart(pubs_by_country)

        log("STEP:docx")
        generate_word_report(
            meta=meta,
            pubs_by_year=pubs_by_year,
            total_docs=total_docs,
            retracted_scopus=retracted,
            retracted_wos=retracted_wos,
            predatory_hits=predatory_hits,
            year_chart_buf=chart_buf,
            country_chart_buf=country_chart_buf,
            pubs_by_country=pubs_by_country,
            output_file=output_file,
            nro_informe=nro_informe,
            institucion=institucion,
        )

        log(f'RESULT:{json.dumps({"output_file": output_file})}')

    # ══════════════════════════════════════════════════════════════════════════
    #  MODE: FULL — Original behavior (collect + generate in one shot)
    # ══════════════════════════════════════════════════════════════════════════
    else:
        nro_informe = params.get("nro_informe", "[COMPLETAR: numero de informe]")
        institucion = params.get("institucion", [
            "Área de Inteligencia e Integridad",
            "Instituto de Investigación Científica",
            "Universidad de Lima",
        ])
        output_file = params.get("output_file", "informe.docx")
        end_year    = date.today().year
        start_year  = end_year - 8

        scopus_id = params.get("scopus_id", "").strip()
        eissn     = params.get("eissn", "").strip()
        article_info = None

        if scopus_id:
            log("STEP:article")
            log(f"INFO:Resolviendo Scopus ID: {scopus_id}")
            article_info = get_article_and_journal_from_scopus_id(api_key, scopus_id)
            log(f"INFO:Articulo: {article_info['article_title'][:80]}")
            log(f"INFO:Revista detectada: {article_info['journal_name']}")
            eissn = article_info["issn_for_lookup"]
            if not eissn:
                raise RuntimeError("No se pudo obtener ISSN")
            log(f"INFO:ISSN para busqueda: {eissn}")

        if not eissn:
            raise RuntimeError("Debes proporcionar 'scopus_id' o 'eissn'")

        log("STEP:metadata")
        meta = get_journal_metadata(api_key, eissn)
        if article_info:
            meta["article_title"]  = article_info["article_title"]
            meta["article_doi"]    = article_info["article_doi"]
            meta["article_date"]   = article_info["article_date"]
            meta["article_author"] = article_info["article_author"]
            meta["scopus_id"]      = scopus_id
            if not meta.get("source_id") and article_info.get("source_id"):
                meta["source_id"] = article_info["source_id"]
        log(f"INFO:Revista: {meta['journal']}")

        log("STEP:total_docs")
        total_docs = get_total_documents(api_key, meta["source_id"])
        log(f"INFO:{total_docs:,} documentos totales")

        log("STEP:pubs_by_year")
        pubs_by_year = get_publications_by_year(api_key, meta["source_id"], start_year, end_year)
        log(f"INFO:{len(pubs_by_year)} anos con datos")

        log("STEP:retracted")
        retracted = get_retracted_count(api_key, meta["source_id"])
        log(f"INFO:{retracted} articulos retractados en Scopus")

        log("STEP:wos_collections")
        wos_api_key  = params.get("wos_api_key", "").strip()
        _init_wos_api(wos_api_key)
        eissn_val_f = meta.get("eissn") or ""
        issn_print_val_f = meta.get("issn_print") or ""
        wos_collections = get_wos_collections(eissn_val_f, issn_print_val_f)
        wos_categories = get_wos_categories(eissn_val_f, issn_print_val_f)
        meta["wos_collections"] = wos_collections
        meta["wos_categories"] = wos_categories
        if wos_collections:
            log(f"INFO:WoS colecciones: {wos_collections}")
        else:
            log("WARN:No se encontraron colecciones WoS")

        log("STEP:wos_retracted")
        retracted_wos = get_wos_retracted_count(eissn_val_f, issn_print_val_f)
        if retracted_wos >= 0:
            log(f"INFO:Retractados WoS: {retracted_wos}")
        else:
            log("WARN:No se pudo obtener retractados WoS")

        log("STEP:predatory")
        hits = check_predatory(meta["journal"], meta["publisher"])
        if hits:
            log(f"WARN:DEPREDADORA detectada: {hits}")
        else:
            log("INFO:No figura en listas depredadoras")

        log("STEP:chart")
        chart_buf = generate_year_chart(pubs_by_year)

        log("STEP:countries")
        issn_for_country = meta.get("eissn") or meta.get("issn_print") or ""
        pubs_by_country = get_publications_by_country(api_key, issn_for_country, top_n=15)
        log(f"INFO:{len(pubs_by_country)} paises con datos")

        log("STEP:country_chart")
        country_chart_buf = generate_country_chart(pubs_by_country)

        log("STEP:docx")
        generate_word_report(
            meta=meta,
            pubs_by_year=pubs_by_year,
            total_docs=total_docs,
            retracted_scopus=retracted,
            retracted_wos=retracted_wos,
            predatory_hits=hits,
            year_chart_buf=chart_buf,
            country_chart_buf=country_chart_buf,
            pubs_by_country=pubs_by_country,
            output_file=output_file,
            nro_informe=nro_informe,
            institucion=institucion,
        )

        result = {
            "journal":     meta["journal"],
            "publisher":   meta["publisher"],
            "total_docs":  total_docs,
            "retracted":   retracted,
            "predatory":   hits,
            "output_file": output_file,
        }
        log(f"RESULT:{json.dumps(result, ensure_ascii=False)}")
